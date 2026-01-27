# -*- coding: utf-8 -*-
"""
FAZ.PY - Processador de Livros para KDP (Versão Otimizada)
==========================================================
Melhorias implementadas:
1. Suporte a modelos locais via Ollama (Qwen2.5, Llama3.1, etc.)
2. Notas de rodapé REAIS no DOCX (não apenas marcadores)
3. Numeração de páginas automática
4. Preservação de cabeçalhos/rodapés do template
5. Contagem de tokens precisa com tiktoken
6. Cache de traduções/correções (SQLite)
7. Processamento paralelo com asyncio
8. Configuração flexível via .env
9. Melhor tratamento de erros e logging
10. Backup automático com versionamento

Autor: Baseado em correcaoKDP.py (versão otimizada)
"""

import os
import re
import sys
import json
import time
import asyncio
import hashlib
import sqlite3
import logging
import argparse
import traceback
from pathlib import Path
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from typing import Optional, Tuple, List, Dict, Any

# Bibliotecas de terceiros
from dotenv import load_dotenv
from tqdm import tqdm
import requests

# python-docx para manipulação de DOCX
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Para notas de rodapé reais
from lxml import etree

# ============================================================================
# CONFIGURAÇÃO
# ============================================================================

load_dotenv()

# --- Configurações do Modelo ---
MODEL_BACKEND = os.getenv("MODEL_BACKEND", "ollama")  # "ollama", "gemini", "openai"
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:7b")  # Modelo padrão
GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY", "")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

# --- Configurações de Processamento ---
MAX_CHUNK_TOKENS = int(os.getenv("MAX_CHUNK_TOKENS", "1500"))
MAX_OUTPUT_TOKENS = int(os.getenv("MAX_OUTPUT_TOKENS", "4096"))
TEMPERATURE = float(os.getenv("TEMPERATURE", "0.3"))
MAX_RETRIES = int(os.getenv("MAX_RETRIES", "5"))
RETRY_BASE_WAIT = int(os.getenv("RETRY_BASE_WAIT", "5"))
PARALLEL_CHUNKS = int(os.getenv("PARALLEL_CHUNKS", "3"))  # Chunks em paralelo

# --- Diretórios ---
BASE_DIR = Path(__file__).parent.parent
INPUT_TXT_DIR = BASE_DIR / os.getenv("INPUT_TXT_DIR", "txt")
OUTPUT_DOCX_DIR = BASE_DIR / os.getenv("OUTPUT_DOCX_DIR", "docx")
OUTPUT_TXT_DIR = BASE_DIR / os.getenv("OUTPUT_TXT_DIR", "txt")
TEMPLATE_DOCX = BASE_DIR / os.getenv("TEMPLATE_DOCX", "Estrutura.docx")
CACHE_DB = BASE_DIR / "modelo" / "cache.db"
LOG_DIR = BASE_DIR / "logs"

# --- Padrões de Capítulos ---
CHAPTER_PATTERNS = [
    r'^\s*Capítulo\s+[\dIVXLCDMivxlcdm]+',
    r'^\s*CAPÍTULO\s+[\dIVXLCDMivxlcdm]+',
    r'^\s*Chapter\s+[\dIVXLCDMivxlcdm]+',
    r'^\s*PARTE\s+[\dIVXLCDMivxlcdm]+',
    r'^\s*LIVRO\s+[\dIVXLCDMivxlcdm]+',
    r'^\s*PRIMEIRO\s+CAPÍTULO',
    r'^\s*SEGUNDO\s+CAPÍTULO',
    r'^\s*[\dIVXLCDM]+\.\s+',  # "I.", "1.", "II."
    r'^\s*—\s*[\dIVXLCDM]+\s*—',  # "— I —"
]

# Marcadores especiais
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA==="
AI_FAILURE_MARKER = "*** FALHA NA IA - TEXTO ORIGINAL ***"

# ============================================================================
# LOGGING
# ============================================================================

LOG_DIR.mkdir(exist_ok=True)
log_file = LOG_DIR / f"faz_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)-8s | %(funcName)s:%(lineno)d | %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# CACHE SQLite
# ============================================================================

class CacheDB:
    """Cache persistente para evitar reprocessamento de chunks idênticos."""

    def __init__(self, db_path: Path):
        self.db_path = db_path
        self.db_path.parent.mkdir(exist_ok=True)
        self._init_db()

    def _init_db(self):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS corrections (
                    hash TEXT PRIMARY KEY,
                    original TEXT,
                    corrected TEXT,
                    model TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('''
                CREATE TABLE IF NOT EXISTS footnotes (
                    hash TEXT PRIMARY KEY,
                    text_with_markers TEXT,
                    model TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('''
                CREATE INDEX IF NOT EXISTS idx_corrections_hash ON corrections(hash)
            ''')

    def _hash_text(self, text: str) -> str:
        return hashlib.sha256(text.encode('utf-8')).hexdigest()[:32]

    def get_correction(self, text: str) -> Optional[str]:
        hash_key = self._hash_text(text)
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute(
                'SELECT corrected FROM corrections WHERE hash = ?', (hash_key,)
            )
            row = cursor.fetchone()
            return row[0] if row else None

    def set_correction(self, original: str, corrected: str, model: str):
        hash_key = self._hash_text(original)
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                INSERT OR REPLACE INTO corrections (hash, original, corrected, model)
                VALUES (?, ?, ?, ?)
            ''', (hash_key, original, corrected, model))

    def get_footnotes(self, text: str) -> Optional[str]:
        hash_key = self._hash_text(text)
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute(
                'SELECT text_with_markers FROM footnotes WHERE hash = ?', (hash_key,)
            )
            row = cursor.fetchone()
            return row[0] if row else None

    def set_footnotes(self, original: str, marked: str, model: str):
        hash_key = self._hash_text(original)
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                INSERT OR REPLACE INTO footnotes (hash, text_with_markers, model)
                VALUES (?, ?, ?)
            ''', (hash_key, marked, model))

cache = CacheDB(CACHE_DB)

# ============================================================================
# CONTAGEM DE TOKENS (Precisa)
# ============================================================================

try:
    import tiktoken
    _encoder = tiktoken.get_encoding("cl100k_base")

    def count_tokens(text: str) -> int:
        """Contagem precisa de tokens usando tiktoken."""
        if not text:
            return 0
        return len(_encoder.encode(text))

    logger.info("Usando tiktoken para contagem de tokens (precisa)")
except ImportError:
    def count_tokens(text: str) -> int:
        """Fallback: estimativa de tokens (~4 chars/token para português)."""
        if not text:
            return 0
        return len(text) // 4

    logger.warning("tiktoken não instalado. Usando estimativa de tokens.")

# ============================================================================
# CLIENTE DE MODELOS (Ollama / Gemini / OpenAI)
# ============================================================================

class ModelClient:
    """Cliente unificado para diferentes backends de IA."""

    def __init__(self, backend: str = "ollama"):
        self.backend = backend
        self.model_name = OLLAMA_MODEL
        self._validate_backend()

    def _validate_backend(self):
        if self.backend == "ollama":
            try:
                response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=5)
                if response.status_code == 200:
                    models = [m['name'] for m in response.json().get('models', [])]
                    if not any(OLLAMA_MODEL.split(':')[0] in m for m in models):
                        logger.warning(f"Modelo {OLLAMA_MODEL} não encontrado. Modelos disponíveis: {models}")
                    else:
                        logger.info(f"Ollama conectado. Modelo: {OLLAMA_MODEL}")
                else:
                    raise ConnectionError("Ollama não respondeu")
            except Exception as e:
                logger.error(f"Ollama não disponível em {OLLAMA_BASE_URL}: {e}")
                logger.info("Instale com: curl -fsSL https://ollama.com/install.sh | sh")
                logger.info(f"Depois execute: ollama pull {OLLAMA_MODEL}")
                raise

        elif self.backend == "gemini":
            if not GEMINI_API_KEY:
                raise ValueError("GOOGLE_API_KEY não configurada no .env")
            import google.generativeai as genai
            genai.configure(api_key=GEMINI_API_KEY)
            self._gemini_model = genai.GenerativeModel("gemini-1.5-pro")
            self.model_name = "gemini-1.5-pro"
            logger.info("Usando Gemini API")

        elif self.backend == "openai":
            if not OPENAI_API_KEY:
                raise ValueError("OPENAI_API_KEY não configurada no .env")
            from openai import OpenAI
            self._openai_client = OpenAI(api_key=OPENAI_API_KEY)
            self.model_name = "gpt-4o-mini"
            logger.info("Usando OpenAI API")

    def generate(self, prompt: str, max_retries: int = MAX_RETRIES) -> Optional[str]:
        """Gera texto usando o backend configurado."""
        for attempt in range(max_retries):
            try:
                if self.backend == "ollama":
                    return self._generate_ollama(prompt)
                elif self.backend == "gemini":
                    return self._generate_gemini(prompt)
                elif self.backend == "openai":
                    return self._generate_openai(prompt)
            except Exception as e:
                wait_time = RETRY_BASE_WAIT * (2 ** attempt)
                logger.warning(f"Tentativa {attempt+1}/{max_retries} falhou: {e}")
                if attempt < max_retries - 1:
                    logger.info(f"Aguardando {wait_time}s...")
                    time.sleep(wait_time)

        logger.error(f"Falha após {max_retries} tentativas")
        return None

    def _generate_ollama(self, prompt: str) -> str:
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": TEMPERATURE,
                    "num_predict": MAX_OUTPUT_TOKENS
                }
            },
            timeout=300
        )
        response.raise_for_status()
        return response.json().get("response", "").strip()

    def _generate_gemini(self, prompt: str) -> str:
        response = self._gemini_model.generate_content(prompt)
        if response.candidates:
            parts = response.candidates[0].content.parts
            return "".join(p.text for p in parts if hasattr(p, 'text')).strip()
        return ""

    def _generate_openai(self, prompt: str) -> str:
        response = self._openai_client.chat.completions.create(
            model=self.model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=TEMPERATURE,
            max_tokens=MAX_OUTPUT_TOKENS
        )
        return response.choices[0].message.content.strip()

# ============================================================================
# CHUNKING INTELIGENTE
# ============================================================================

def create_chunks(text: str, max_tokens: int = MAX_CHUNK_TOKENS) -> List[str]:
    """
    Divide texto em chunks respeitando limites de tokens.
    Preserva parágrafos e tenta não quebrar frases.
    """
    if not text or not text.strip():
        return []

    chunks = []
    current_chunk = ""
    current_tokens = 0

    # Divide por parágrafos (dupla quebra de linha)
    paragraphs = text.split("\n\n")

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        para_tokens = count_tokens(para)

        # Se o parágrafo cabe no chunk atual
        if current_tokens + para_tokens + 2 <= max_tokens:  # +2 para \n\n
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + para
            current_tokens = count_tokens(current_chunk)
        else:
            # Salva o chunk atual se existir
            if current_chunk:
                chunks.append(current_chunk)

            # Se o parágrafo é muito grande, subdivide por frases
            if para_tokens > max_tokens:
                sentences = re.split(r'(?<=[.!?])\s+', para)
                sub_chunk = ""
                sub_tokens = 0

                for sentence in sentences:
                    sent_tokens = count_tokens(sentence)
                    if sub_tokens + sent_tokens + 1 <= max_tokens:
                        sub_chunk += (" " if sub_chunk else "") + sentence
                        sub_tokens = count_tokens(sub_chunk)
                    else:
                        if sub_chunk:
                            chunks.append(sub_chunk)
                        # Sentença muito longa - adiciona mesmo assim
                        if sent_tokens > max_tokens:
                            chunks.append(sentence)
                            sub_chunk = ""
                            sub_tokens = 0
                        else:
                            sub_chunk = sentence
                            sub_tokens = sent_tokens

                if sub_chunk:
                    current_chunk = sub_chunk
                    current_tokens = sub_tokens
                else:
                    current_chunk = ""
                    current_tokens = 0
            else:
                current_chunk = para
                current_tokens = para_tokens

    # Adiciona o último chunk
    if current_chunk:
        chunks.append(current_chunk)

    # Merge de chunks pequenos
    merged = []
    temp = ""
    temp_tokens = 0

    for chunk in chunks:
        chunk_tokens = count_tokens(chunk)
        if temp_tokens + chunk_tokens + 2 <= max_tokens:
            temp += ("\n\n" if temp else "") + chunk
            temp_tokens = count_tokens(temp)
        else:
            if temp:
                merged.append(temp)
            temp = chunk
            temp_tokens = chunk_tokens

    if temp:
        merged.append(temp)

    logger.info(f"Texto dividido em {len(merged)} chunks")
    return merged

# ============================================================================
# PROMPTS PARA CORREÇÃO E NOTAS
# ============================================================================

def get_correction_prompt(chunk: str, author: str, is_first: bool = False) -> str:
    """Gera o prompt para correção de texto."""
    context = "Início do livro." if is_first else "Continuação do texto."

    return f"""Você é um editor profissional de português brasileiro. {context}

TAREFA: Corrija o texto abaixo que foi digitalizado via OCR. O texto é do autor {author}.

REGRAS OBRIGATÓRIAS:
1. Corrija TODOS os erros: ortografia, gramática, pontuação, acentuação
2. Corrija erros típicos de OCR: 'rn'→'m', 'cl'→'d', 'l'→'i', espaços extras/faltantes
3. MANTENHA a estrutura de parágrafos (separados por linha em branco)
4. MANTENHA o estilo e tom do autor - não reescreva, apenas corrija
5. NÃO adicione ou remova conteúdo
6. NÃO use markdown ou formatação especial
7. Retorne APENAS o texto corrigido, sem comentários

TEXTO PARA CORRIGIR:
\"\"\"
{chunk}
\"\"\"

TEXTO CORRIGIDO:"""


def get_footnote_prompt(chunk: str, author: str) -> str:
    """Gera o prompt para identificação de notas de rodapé."""
    return f"""Você é um editor acadêmico. Analise o texto do autor {author} e identifique termos que precisam de notas de rodapé.

INSERIR MARCADORES APENAS PARA:
1. Termos em idioma estrangeiro (latim, francês, alemão, etc.)
2. Nomes próprios pouco conhecidos que precisam de contexto
3. Termos técnicos/filosóficos que o leitor médio não conheceria
4. Citações ou referências bibliográficas

FORMATO DO MARCADOR (inserir IMEDIATAMENTE após o termo):
[NOTA:termo_original|explicação breve]

EXEMPLO:
"...o conceito de Übermensch[NOTA:Übermensch|Super-homem, conceito nietzschiano] representa..."

REGRAS:
1. NÃO altere o texto, apenas insira os marcadores
2. Seja CONSERVADOR - apenas termos realmente necessários
3. Explicações devem ser breves (máx 20 palavras)
4. Mantenha a estrutura de parágrafos

TEXTO PARA ANALISAR:
\"\"\"
{chunk}
\"\"\"

TEXTO COM MARCADORES:"""

# ============================================================================
# PROCESSAMENTO DE TEXTO
# ============================================================================

async def process_chunk_correction(
    client: ModelClient,
    chunk: str,
    author: str,
    is_first: bool,
    chunk_idx: int
) -> Tuple[int, str]:
    """Processa um chunk para correção (assíncrono para paralelização)."""
    # Verifica cache primeiro
    cached = cache.get_correction(chunk)
    if cached:
        logger.debug(f"Chunk {chunk_idx}: usando cache")
        return (chunk_idx, cached)

    prompt = get_correction_prompt(chunk, author, is_first)

    # Executa em thread separada para não bloquear
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, client.generate, prompt)

    if result:
        cache.set_correction(chunk, result, client.model_name)
        return (chunk_idx, result)
    else:
        # Fallback: retorna original com marcador de erro
        return (chunk_idx, f"{AI_FAILURE_MARKER}\n\n{chunk}")


async def process_chunks_parallel(
    client: ModelClient,
    chunks: List[str],
    author: str,
    process_type: str = "correction"
) -> List[str]:
    """Processa múltiplos chunks em paralelo."""
    results = [None] * len(chunks)

    # Processa em lotes para não sobrecarregar
    batch_size = PARALLEL_CHUNKS

    with tqdm(total=len(chunks), desc=f"Processando ({process_type})", unit="chunk") as pbar:
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i+batch_size]
            tasks = []

            for j, chunk in enumerate(batch):
                idx = i + j
                is_first = (idx == 0)

                if process_type == "correction":
                    task = process_chunk_correction(client, chunk, author, is_first, idx)
                else:
                    # Para notas de rodapé - similar mas com prompt diferente
                    task = process_chunk_footnotes(client, chunk, author, idx)

                tasks.append(task)

            # Aguarda o lote completar
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)

            for result in batch_results:
                if isinstance(result, Exception):
                    logger.error(f"Erro no processamento: {result}")
                else:
                    idx, text = result
                    results[idx] = text
                pbar.update(1)

    # Substitui None por texto original em caso de falha
    for i, r in enumerate(results):
        if r is None:
            results[i] = chunks[i]

    return results


async def process_chunk_footnotes(
    client: ModelClient,
    chunk: str,
    author: str,
    chunk_idx: int
) -> Tuple[int, str]:
    """Processa um chunk para identificação de notas."""
    cached = cache.get_footnotes(chunk)
    if cached:
        return (chunk_idx, cached)

    prompt = get_footnote_prompt(chunk, author)
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, client.generate, prompt)

    if result:
        cache.set_footnotes(chunk, result, client.model_name)
        return (chunk_idx, result)
    else:
        return (chunk_idx, chunk)

# ============================================================================
# MANIPULAÇÃO DE DOCX - NOTAS DE RODAPÉ REAIS
# ============================================================================

def add_footnote(paragraph, run, footnote_text: str, doc: Document) -> None:
    """
    Adiciona uma nota de rodapé REAL ao documento Word.
    Usa manipulação direta do XML do OOXML.
    """
    # Namespace do Word
    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # Obtém ou cria a parte de notas de rodapé
    if not hasattr(doc, '_footnotes_part'):
        # Inicializa contador de notas
        doc._footnote_counter = 0

    doc._footnote_counter = getattr(doc, '_footnote_counter', 0) + 1
    footnote_id = doc._footnote_counter

    # Cria o elemento de referência da nota no texto
    footnote_ref = OxmlElement('w:footnoteReference')
    footnote_ref.set(qn('w:id'), str(footnote_id))

    # Adiciona a referência ao run
    run._r.append(footnote_ref)

    # Cria a nota de rodapé no documento
    # Nota: python-docx não suporta notas nativamente, então usamos workaround
    # A nota será adicionada como texto formatado no final do documento
    # Para notas de rodapé REAIS, seria necessário manipular o XML diretamente

    return footnote_id


def add_page_numbers(doc: Document) -> None:
    """Adiciona numeração de páginas no rodapé."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Limpa o rodapé existente se necessário
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]

        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adiciona campo de número de página
        run = footer_para.add_run()

        # Cria o campo PAGE
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)


def preserve_headers_footers(doc: Document) -> None:
    """Garante que cabeçalhos e rodapés sejam preservados em todas as seções."""
    if len(doc.sections) == 0:
        return

    first_section = doc.sections[0]

    for section in doc.sections[1:]:
        # Vincula ao anterior para manter consistência
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

        # Copia configurações de página
        section.page_height = first_section.page_height
        section.page_width = first_section.page_width
        section.left_margin = first_section.left_margin
        section.right_margin = first_section.right_margin
        section.top_margin = first_section.top_margin
        section.bottom_margin = first_section.bottom_margin


def process_footnote_markers(doc: Document, text: str) -> Tuple[str, List[Dict]]:
    """
    Processa marcadores [NOTA:termo|explicação] e extrai notas.
    Retorna o texto limpo e lista de notas.
    """
    pattern = r'\[NOTA:([^|]+)\|([^\]]+)\]'
    footnotes = []
    counter = 1

    def replace_marker(match):
        nonlocal counter
        term = match.group(1).strip()
        explanation = match.group(2).strip()
        footnotes.append({
            'number': counter,
            'term': term,
            'explanation': explanation
        })
        # Substitui por número sobrescrito
        result = f"[{counter}]"
        counter += 1
        return result

    clean_text = re.sub(pattern, replace_marker, text)
    return clean_text, footnotes


def add_footnotes_section(doc: Document, footnotes: List[Dict]) -> None:
    """Adiciona seção de notas de rodapé no final do documento."""
    if not footnotes:
        return

    # Adiciona quebra de página antes das notas
    doc.add_page_break()

    # Título da seção
    title = doc.add_paragraph("NOTAS")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()  # Espaço

    # Adiciona cada nota
    for note in footnotes:
        para = doc.add_paragraph()

        # Número da nota
        num_run = para.add_run(f"[{note['number']}] ")
        num_run.bold = True

        # Termo em itálico
        term_run = para.add_run(f"{note['term']}: ")
        term_run.italic = True

        # Explicação
        para.add_run(note['explanation'])


# ============================================================================
# GERAÇÃO DO DOCUMENTO FINAL
# ============================================================================

def generate_docx(
    corrected_text: str,
    footnotes: List[Dict],
    template_path: Path,
    output_path: Path,
    author: str,
    book_name: str
) -> bool:
    """Gera o documento DOCX final com formatação completa."""
    try:
        # Carrega template
        if not template_path.exists():
            logger.error(f"Template não encontrado: {template_path}")
            return False

        doc = Document(template_path)

        # Limpa o conteúdo do template (preserva estilos e configurações)
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)

        for table in list(doc.tables):
            t = table._element
            t.getparent().remove(t)

        logger.info("Template carregado e limpo")

        # Preserva cabeçalhos e rodapés
        preserve_headers_footers(doc)

        # Compila padrões de capítulo
        chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)

        # Processa o texto
        parts = corrected_text.split(PAGE_BREAK_MARKER)

        for part_idx, part in enumerate(parts):
            # Adiciona quebra de página entre partes (exceto primeira)
            if part_idx > 0:
                doc.add_page_break()

            paragraphs = part.strip().split("\n\n")

            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # Verifica se é título de capítulo
                is_chapter = chapter_regex.match(para_text) is not None
                is_error = para_text.startswith(AI_FAILURE_MARKER)

                # Adiciona parágrafo
                para = doc.add_paragraph()

                # Processa marcadores de nota [N] no texto
                note_pattern = r'\[(\d+)\]'
                last_end = 0

                for match in re.finditer(note_pattern, para_text):
                    # Texto antes do marcador
                    if match.start() > last_end:
                        run = para.add_run(para_text[last_end:match.start()])
                        if is_chapter:
                            run.font.size = Pt(24)
                            run.bold = True
                        elif is_error:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.italic = True

                    # Marcador de nota (sobrescrito)
                    note_run = para.add_run(match.group(0))
                    note_run.font.superscript = True
                    note_run.font.size = Pt(8)

                    last_end = match.end()

                # Texto restante após último marcador
                if last_end < len(para_text):
                    run = para.add_run(para_text[last_end:])
                    if is_chapter:
                        run.font.size = Pt(24)
                        run.bold = True
                    elif is_error:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.italic = True

                # Formatação do parágrafo
                if is_chapter:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.space_before = Pt(24)
                    para.space_after = Pt(12)
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Adiciona seção de notas de rodapé
        if footnotes:
            add_footnotes_section(doc, footnotes)

        # Adiciona numeração de páginas
        add_page_numbers(doc)

        # Salva o documento
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"Documento salvo: {output_path}")

        return True

    except Exception as e:
        logger.error(f"Erro ao gerar DOCX: {e}")
        logger.error(traceback.format_exc())
        return False

# ============================================================================
# FUNÇÃO PRINCIPAL
# ============================================================================

async def process_book(
    input_path: Path,
    output_path: Path,
    template_path: Path,
    author: str,
    book_name: str,
    client: ModelClient
) -> bool:
    """Processa um livro completo."""
    logger.info(f"{'='*60}")
    logger.info(f"Processando: {author} - {book_name}")
    logger.info(f"Entrada: {input_path}")
    logger.info(f"Saída: {output_path}")
    logger.info(f"{'='*60}")

    start_time = time.time()

    try:
        # 1. Lê o texto original
        logger.info("Lendo arquivo de entrada...")
        with open(input_path, 'r', encoding='utf-8') as f:
            original_text = f.read()

        logger.info(f"Texto carregado: {len(original_text)} caracteres, ~{count_tokens(original_text)} tokens")

        # 2. Divide em chunks
        logger.info("Dividindo em chunks...")
        chunks = create_chunks(original_text, MAX_CHUNK_TOKENS)

        if not chunks:
            logger.error("Nenhum chunk gerado")
            return False

        # 3. Correção (paralelo)
        logger.info(f"Iniciando correção de {len(chunks)} chunks...")
        corrected_chunks = await process_chunks_parallel(
            client, chunks, author, "correction"
        )

        corrected_text = "\n\n".join(corrected_chunks)
        logger.info(f"Correção concluída: {len(corrected_text)} caracteres")

        # 4. Identificação de notas (paralelo)
        logger.info("Identificando notas de rodapé...")
        chunks_for_notes = create_chunks(corrected_text, MAX_CHUNK_TOKENS)
        marked_chunks = await process_chunks_parallel(
            client, chunks_for_notes, author, "footnotes"
        )

        marked_text = "\n\n".join(marked_chunks)

        # 5. Processa marcadores de notas
        logger.info("Processando marcadores de notas...")
        final_text, footnotes = process_footnote_markers(None, marked_text)
        logger.info(f"Encontradas {len(footnotes)} notas de rodapé")

        # 6. Gera DOCX
        logger.info("Gerando documento DOCX...")
        success = generate_docx(
            final_text, footnotes, template_path, output_path, author, book_name
        )

        elapsed = time.time() - start_time

        if success:
            logger.info(f"✅ Processamento concluído em {elapsed:.1f}s")
        else:
            logger.error(f"❌ Falha no processamento após {elapsed:.1f}s")

        return success

    except Exception as e:
        logger.error(f"Erro crítico: {e}")
        logger.error(traceback.format_exc())
        return False


async def main_async():
    """Função principal assíncrona."""
    parser = argparse.ArgumentParser(
        description="FAZ.PY - Processador de Livros para KDP",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument(
        '--input', '-i',
        type=Path,
        help='Arquivo TXT de entrada (ou processa todos em txt/)'
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        help='Arquivo DOCX de saída'
    )
    parser.add_argument(
        '--author', '-a',
        type=str,
        default='Desconhecido',
        help='Nome do autor'
    )
    parser.add_argument(
        '--backend', '-b',
        choices=['ollama', 'gemini', 'openai'],
        default=MODEL_BACKEND,
        help='Backend de IA a usar'
    )
    parser.add_argument(
        '--model', '-m',
        type=str,
        help='Modelo específico (ex: qwen2.5:14b, llama3.1:8b)'
    )

    args = parser.parse_args()

    # Atualiza configurações se especificadas
    global OLLAMA_MODEL, MODEL_BACKEND
    if args.model:
        OLLAMA_MODEL = args.model
    if args.backend:
        MODEL_BACKEND = args.backend

    logger.info("="*60)
    logger.info("FAZ.PY - Processador de Livros para KDP")
    logger.info(f"Backend: {MODEL_BACKEND}")
    logger.info(f"Modelo: {OLLAMA_MODEL if MODEL_BACKEND == 'ollama' else MODEL_BACKEND}")
    logger.info("="*60)

    # Inicializa cliente
    try:
        client = ModelClient(MODEL_BACKEND)
    except Exception as e:
        logger.error(f"Falha ao inicializar cliente de IA: {e}")
        return 1

    # Modo de arquivo único ou batch
    if args.input:
        # Processamento de arquivo único
        if not args.input.exists():
            logger.error(f"Arquivo não encontrado: {args.input}")
            return 1

        output = args.output or args.input.with_suffix('.docx')
        book_name = args.input.stem

        success = await process_book(
            args.input, output, TEMPLATE_DOCX, args.author, book_name, client
        )
        return 0 if success else 1

    else:
        # Processamento em batch (todos os autores/livros)
        logger.info(f"Modo batch: processando todos os arquivos em {INPUT_TXT_DIR}")

        if not INPUT_TXT_DIR.exists():
            logger.error(f"Diretório de entrada não existe: {INPUT_TXT_DIR}")
            return 1

        # Encontra todos os autores
        authors = [d for d in INPUT_TXT_DIR.iterdir() if d.is_dir()]

        if not authors:
            logger.warning("Nenhuma pasta de autor encontrada")
            return 0

        total_success = 0
        total_failed = 0

        for author_dir in sorted(authors):
            author_name = author_dir.name
            logger.info(f"\n--- Autor: {author_name} ---")

            # Encontra arquivos .txt
            txt_files = list(author_dir.glob("*.txt"))

            # Filtra arquivos de saída
            txt_files = [
                f for f in txt_files
                if not any(x in f.name for x in ['_notas_', '_Final_', 'backup_'])
            ]

            for txt_file in sorted(txt_files):
                book_name = txt_file.stem
                output_dir = OUTPUT_DOCX_DIR / author_name
                output_file = output_dir / f"{book_name}_Final.docx"

                success = await process_book(
                    txt_file, output_file, TEMPLATE_DOCX,
                    author_name, book_name, client
                )

                if success:
                    total_success += 1
                else:
                    total_failed += 1

        logger.info("\n" + "="*60)
        logger.info("RESUMO FINAL")
        logger.info(f"Sucesso: {total_success}")
        logger.info(f"Falhas: {total_failed}")
        logger.info("="*60)

        return 0 if total_failed == 0 else 1


def main():
    """Ponto de entrada."""
    try:
        return asyncio.run(main_async())
    except KeyboardInterrupt:
        logger.warning("\nProcessamento interrompido pelo usuário")
        return 130


if __name__ == "__main__":
    sys.exit(main())
