# -*- coding: utf-8 -*-
"""
TRADUTOR.PY - Sistema Inteligente de Tradução de Livros
=======================================================
Sistema ROBUSTO e ADAPTÁVEL que:
1. Detecta automaticamente o formato do arquivo (PDF, EPUB, HTML, TXT, DOCX)
2. Extrai texto de forma inteligente com múltiplos fallbacks
3. Detecta idioma automaticamente
4. Traduz para português brasileiro
5. Corrige e formata o texto
6. Gera DOCX profissional
7. Loga erros e continua com próximo livro

RESILIENTE: Se um arquivo falhar, loga o erro e continua!
"""

import os
import re
import sys
import json
import time
import asyncio
import hashlib
import sqlite3
import logging
import traceback
from pathlib import Path
from datetime import datetime
from typing import Optional, Tuple, List, Dict, Any
from dataclasses import dataclass
from enum import Enum

# Bibliotecas de terceiros
from dotenv import load_dotenv
from tqdm import tqdm
import requests

# DOCX
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================================
# IMPORTAÇÕES OPCIONAIS (com fallbacks)
# ============================================================================

# PDF - PyMuPDF
try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# PDF - pdfplumber
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# EPUB - ebooklib
try:
    from ebooklib import epub
    HAS_EBOOKLIB = True
except ImportError:
    HAS_EBOOKLIB = False

# HTML parsing
from html.parser import HTMLParser
import zipfile

# Tokens
try:
    import tiktoken
    _enc = tiktoken.get_encoding("cl100k_base")
    def count_tokens(text: str) -> int:
        return len(_enc.encode(text)) if text else 0
except ImportError:
    def count_tokens(text: str) -> int:
        return len(text) // 4 if text else 0

# ============================================================================
# CONFIGURAÇÃO
# ============================================================================

load_dotenv()

BASE_DIR = Path(__file__).parent.parent
CALIBRE_DIR = Path(os.getenv("CALIBRE_DIR", BASE_DIR.parent / "CALIBRE"))
INPUT_TXT_DIR = BASE_DIR / "txt"
OUTPUT_DOCX_DIR = BASE_DIR / "docx_traduzidos"
TEMPLATE_DOCX = BASE_DIR / "Estrutura.docx"
CACHE_DB = BASE_DIR / "modelo" / "cache.db"
LOG_DIR = BASE_DIR / "logs"

OLLAMA_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:7b")

MAX_CHUNK_TOKENS = 1200
MAX_RETRIES = 3
MIN_TEXT_LENGTH = 500  # Mínimo de caracteres para considerar válido

# ============================================================================
# LOGGING INTELIGENTE
# ============================================================================

LOG_DIR.mkdir(exist_ok=True)

# Log principal
log_file = LOG_DIR / f"tradutor_{datetime.now().strftime('%Y%m%d')}.log"
# Log de erros separado
error_log = LOG_DIR / "erros_extracao.log"
# Log de livros processados
processed_log = LOG_DIR / "livros_processados.log"
# Log de livros com erro
failed_log = LOG_DIR / "livros_falhos.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)-8s | %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# ENUMS E DATACLASSES
# ============================================================================

class FileFormat(Enum):
    PDF = "pdf"
    EPUB = "epub"
    HTML = "html"
    TXT = "txt"
    DOCX = "docx"
    UNKNOWN = "unknown"

class ExtractionStatus(Enum):
    SUCCESS = "success"
    EMPTY = "empty"
    CORRUPTED = "corrupted"
    UNSUPPORTED = "unsupported"
    ERROR = "error"

@dataclass
class ExtractionResult:
    status: ExtractionStatus
    text: Optional[str]
    format: FileFormat
    error_message: Optional[str] = None
    pages_extracted: int = 0

@dataclass
class BookInfo:
    path: Path
    title: str
    author: str
    language: str
    format: FileFormat

# ============================================================================
# GERENCIAMENTO DE LOGS
# ============================================================================

def load_processed_books() -> set:
    """Carrega lista de livros já processados."""
    if not processed_log.exists():
        return set()
    with open(processed_log, 'r', encoding='utf-8') as f:
        return set(line.strip() for line in f if line.strip())

def load_failed_books() -> set:
    """Carrega lista de livros que falharam."""
    if not failed_log.exists():
        return set()
    with open(failed_log, 'r', encoding='utf-8') as f:
        return set(line.strip() for line in f if line.strip())

def mark_as_processed(book_id: str):
    """Marca livro como processado."""
    with open(processed_log, 'a', encoding='utf-8') as f:
        f.write(f"{book_id}\n")

def mark_as_failed(book_id: str, reason: str):
    """Marca livro como falho com razão."""
    with open(failed_log, 'a', encoding='utf-8') as f:
        f.write(f"{book_id}|{reason}|{datetime.now().isoformat()}\n")

def log_extraction_error(file_path: Path, error: str):
    """Loga erro de extração em arquivo separado."""
    with open(error_log, 'a', encoding='utf-8') as f:
        f.write(f"{datetime.now().isoformat()}|{file_path}|{error}\n")

# ============================================================================
# CACHE SQLITE
# ============================================================================

class TranslationCache:
    def __init__(self, db_path: Path):
        self.db_path = db_path
        db_path.parent.mkdir(exist_ok=True)
        self._init_db()

    def _init_db(self):
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute('''
                    CREATE TABLE IF NOT EXISTS translations (
                        hash TEXT PRIMARY KEY,
                        original TEXT,
                        translated TEXT,
                        source_lang TEXT,
                        model TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                ''')
        except Exception as e:
            logger.warning(f"Erro ao inicializar cache: {e}")

    def _hash(self, text: str) -> str:
        return hashlib.sha256(text.encode()).hexdigest()[:32]

    def get(self, text: str) -> Optional[str]:
        try:
            with sqlite3.connect(self.db_path) as conn:
                cur = conn.execute(
                    'SELECT translated FROM translations WHERE hash = ?',
                    (self._hash(text),)
                )
                row = cur.fetchone()
                return row[0] if row else None
        except:
            return None

    def set(self, original: str, translated: str, source_lang: str, model: str):
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute('''
                    INSERT OR REPLACE INTO translations
                    (hash, original, translated, source_lang, model)
                    VALUES (?, ?, ?, ?, ?)
                ''', (self._hash(original), original, translated, source_lang, model))
        except Exception as e:
            logger.warning(f"Erro ao salvar cache: {e}")

cache = TranslationCache(CACHE_DB)

# ============================================================================
# EXTRAÇÃO DE TEXTO (INTELIGENTE COM FALLBACKS)
# ============================================================================

class HTMLTextExtractor(HTMLParser):
    """Parser HTML simples."""
    def __init__(self):
        super().__init__()
        self.text_parts = []
        self.skip_tags = {'script', 'style', 'head', 'meta', 'link', 'nav', 'footer'}
        self.current_tag = None

    def handle_starttag(self, tag, attrs):
        self.current_tag = tag.lower()
        if tag.lower() in ('p', 'div', 'br', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'):
            self.text_parts.append('\n\n')

    def handle_data(self, data):
        if self.current_tag not in self.skip_tags:
            text = data.strip()
            if text:
                self.text_parts.append(text + ' ')

    def get_text(self) -> str:
        text = ''.join(self.text_parts)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        return text.strip()


def detect_format(file_path: Path) -> FileFormat:
    """Detecta formato do arquivo."""
    suffix = file_path.suffix.lower()
    format_map = {
        '.pdf': FileFormat.PDF,
        '.epub': FileFormat.EPUB,
        '.html': FileFormat.HTML,
        '.htm': FileFormat.HTML,
        '.xhtml': FileFormat.HTML,
        '.txt': FileFormat.TXT,
        '.docx': FileFormat.DOCX,
    }
    return format_map.get(suffix, FileFormat.UNKNOWN)


def extract_from_pdf(file_path: Path) -> ExtractionResult:
    """Extrai texto de PDF com múltiplos métodos."""
    texts = []
    pages = 0

    # Método 1: PyMuPDF (preferido)
    if HAS_PYMUPDF:
        try:
            doc = fitz.open(str(file_path))
            for page in doc:
                text = page.get_text()
                if text and text.strip():
                    texts.append(text.strip())
                    pages += 1
            doc.close()

            if texts:
                return ExtractionResult(
                    status=ExtractionStatus.SUCCESS,
                    text='\n\n'.join(texts),
                    format=FileFormat.PDF,
                    pages_extracted=pages
                )
        except Exception as e:
            logger.debug(f"PyMuPDF falhou: {e}")

    # Método 2: pdfplumber (fallback)
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append(text.strip())
                        pages += 1

            if texts:
                return ExtractionResult(
                    status=ExtractionStatus.SUCCESS,
                    text='\n\n'.join(texts),
                    format=FileFormat.PDF,
                    pages_extracted=pages
                )
        except Exception as e:
            logger.debug(f"pdfplumber falhou: {e}")

    # Sem bibliotecas PDF
    if not HAS_PYMUPDF and not HAS_PDFPLUMBER:
        return ExtractionResult(
            status=ExtractionStatus.UNSUPPORTED,
            text=None,
            format=FileFormat.PDF,
            error_message="Nenhuma biblioteca PDF instalada (pip install pymupdf pdfplumber)"
        )

    return ExtractionResult(
        status=ExtractionStatus.CORRUPTED,
        text=None,
        format=FileFormat.PDF,
        error_message="PDF corrompido ou protegido"
    )


def extract_from_epub(file_path: Path) -> ExtractionResult:
    """Extrai texto de EPUB."""
    texts = []

    # Método 1: ebooklib
    if HAS_EBOOKLIB:
        try:
            book = epub.read_epub(str(file_path))
            for item in book.get_items():
                if item.get_type() == 9:  # ITEM_DOCUMENT
                    content = item.get_content().decode('utf-8', errors='ignore')
                    parser = HTMLTextExtractor()
                    parser.feed(content)
                    text = parser.get_text()
                    if text and len(text) > 50:
                        texts.append(text)

            if texts:
                return ExtractionResult(
                    status=ExtractionStatus.SUCCESS,
                    text='\n\n'.join(texts),
                    format=FileFormat.EPUB
                )
        except Exception as e:
            logger.debug(f"ebooklib falhou: {e}")

    # Método 2: Extração manual do ZIP
    try:
        with zipfile.ZipFile(str(file_path), 'r') as zf:
            for name in zf.namelist():
                if name.endswith(('.html', '.xhtml', '.htm')):
                    try:
                        content = zf.read(name).decode('utf-8', errors='ignore')
                        parser = HTMLTextExtractor()
                        parser.feed(content)
                        text = parser.get_text()
                        if text and len(text) > 50:
                            texts.append(text)
                    except:
                        continue

        if texts:
            return ExtractionResult(
                status=ExtractionStatus.SUCCESS,
                text='\n\n'.join(texts),
                format=FileFormat.EPUB
            )
    except Exception as e:
        logger.debug(f"Extração ZIP falhou: {e}")

    return ExtractionResult(
        status=ExtractionStatus.CORRUPTED,
        text=None,
        format=FileFormat.EPUB,
        error_message="EPUB corrompido ou formato inválido"
    )


def extract_from_html(file_path: Path) -> ExtractionResult:
    """Extrai texto de HTML."""
    try:
        # Tenta diferentes encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            # Fallback: ignora erros
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

        parser = HTMLTextExtractor()
        parser.feed(content)
        text = parser.get_text()

        if text and len(text) > MIN_TEXT_LENGTH:
            return ExtractionResult(
                status=ExtractionStatus.SUCCESS,
                text=text,
                format=FileFormat.HTML
            )
        else:
            return ExtractionResult(
                status=ExtractionStatus.EMPTY,
                text=None,
                format=FileFormat.HTML,
                error_message="HTML sem conteúdo útil"
            )

    except Exception as e:
        return ExtractionResult(
            status=ExtractionStatus.ERROR,
            text=None,
            format=FileFormat.HTML,
            error_message=str(e)
        )


def extract_from_txt(file_path: Path) -> ExtractionResult:
    """Extrai texto de TXT."""
    try:
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

        if text and len(text.strip()) > MIN_TEXT_LENGTH:
            return ExtractionResult(
                status=ExtractionStatus.SUCCESS,
                text=text.strip(),
                format=FileFormat.TXT
            )
        else:
            return ExtractionResult(
                status=ExtractionStatus.EMPTY,
                text=None,
                format=FileFormat.TXT,
                error_message="Arquivo muito pequeno ou vazio"
            )

    except Exception as e:
        return ExtractionResult(
            status=ExtractionStatus.ERROR,
            text=None,
            format=FileFormat.TXT,
            error_message=str(e)
        )


def extract_from_docx(file_path: Path) -> ExtractionResult:
    """Extrai texto de DOCX."""
    try:
        doc = Document(str(file_path))
        texts = [para.text for para in doc.paragraphs if para.text.strip()]
        text = '\n\n'.join(texts)

        if text and len(text) > MIN_TEXT_LENGTH:
            return ExtractionResult(
                status=ExtractionStatus.SUCCESS,
                text=text,
                format=FileFormat.DOCX
            )
        else:
            return ExtractionResult(
                status=ExtractionStatus.EMPTY,
                text=None,
                format=FileFormat.DOCX,
                error_message="DOCX sem conteúdo útil"
            )

    except Exception as e:
        return ExtractionResult(
            status=ExtractionStatus.CORRUPTED,
            text=None,
            format=FileFormat.DOCX,
            error_message=str(e)
        )


def extract_text_smart(file_path: Path) -> ExtractionResult:
    """
    EXTRAÇÃO INTELIGENTE - Detecta formato e extrai com fallbacks.
    NUNCA falha silenciosamente - sempre retorna um resultado.
    """
    logger.info(f"Extraindo: {file_path.name}")

    if not file_path.exists():
        return ExtractionResult(
            status=ExtractionStatus.ERROR,
            text=None,
            format=FileFormat.UNKNOWN,
            error_message="Arquivo não encontrado"
        )

    file_format = detect_format(file_path)
    logger.info(f"Formato detectado: {file_format.value}")

    extractors = {
        FileFormat.PDF: extract_from_pdf,
        FileFormat.EPUB: extract_from_epub,
        FileFormat.HTML: extract_from_html,
        FileFormat.TXT: extract_from_txt,
        FileFormat.DOCX: extract_from_docx,
    }

    extractor = extractors.get(file_format)

    if extractor:
        try:
            result = extractor(file_path)
            logger.info(f"Status: {result.status.value}")
            return result
        except Exception as e:
            error_msg = f"Erro inesperado: {str(e)}"
            logger.error(error_msg)
            log_extraction_error(file_path, error_msg)
            return ExtractionResult(
                status=ExtractionStatus.ERROR,
                text=None,
                format=file_format,
                error_message=error_msg
            )

    # Formato desconhecido - tenta como texto
    logger.warning(f"Formato desconhecido, tentando como texto")
    return extract_from_txt(file_path)

# ============================================================================
# DETECÇÃO DE IDIOMA
# ============================================================================

def detect_language(text: str) -> str:
    """Detecta idioma do texto."""
    sample = text[:5000].lower()

    patterns = {
        'pt': ['não', 'você', 'está', 'também', 'porque', 'quando', 'muito', 'então', 'mesmo', 'todos', 'tinha', 'seria', 'fazer', 'ainda', 'depois'],
        'es': ['que', 'para', 'pero', 'más', 'cuando', 'también', 'puede', 'hacer', 'después', 'entonces', 'mismo', 'todos', 'había', 'porque', 'donde'],
        'en': ['the', 'and', 'that', 'have', 'for', 'not', 'with', 'you', 'this', 'but', 'from', 'they', 'would', 'there', 'their', 'what', 'about'],
        'fr': ['les', 'des', 'une', 'que', 'dans', 'pour', 'qui', 'sur', 'est', 'pas', 'plus', 'avec', 'son', 'sont', 'mais'],
        'de': ['und', 'der', 'die', 'das', 'ist', 'nicht', 'von', 'sie', 'mit', 'sich', 'auf', 'für', 'als', 'auch', 'aber'],
    }

    # Caracteres específicos
    char_hints = {
        'pt': ['ç', 'ã', 'õ'],
        'es': ['ñ', '¿', '¡'],
        'de': ['ß', 'ü', 'ö', 'ä'],
        'fr': ['œ', 'æ', 'ç'],
    }

    scores = {lang: 0 for lang in patterns}

    for lang, words in patterns.items():
        for word in words:
            if f' {word} ' in sample:
                scores[lang] += 1

    for lang, chars in char_hints.items():
        for char in chars:
            if char in sample:
                scores[lang] += 3

    best = max(scores, key=scores.get)
    return best if scores[best] >= 3 else 'unknown'

# ============================================================================
# CLIENTE OLLAMA
# ============================================================================

class OllamaClient:
    def __init__(self, model: str = OLLAMA_MODEL):
        self.model = model
        self.base_url = OLLAMA_URL
        self._verify_connection()

    def _verify_connection(self):
        try:
            r = requests.get(f"{self.base_url}/api/tags", timeout=10)
            if r.status_code == 200:
                models = [m['name'] for m in r.json().get('models', [])]
                if self.model not in models and self.model.split(':')[0] not in [m.split(':')[0] for m in models]:
                    logger.warning(f"Modelo {self.model} não encontrado. Disponíveis: {models}")
                else:
                    logger.info(f"✓ Ollama conectado. Modelo: {self.model}")
        except Exception as e:
            logger.error(f"✗ Ollama não disponível: {e}")
            raise ConnectionError("Ollama não está rodando")

    def generate(self, prompt: str) -> Optional[str]:
        for attempt in range(MAX_RETRIES):
            try:
                r = requests.post(
                    f"{self.base_url}/api/generate",
                    json={
                        "model": self.model,
                        "prompt": prompt,
                        "stream": False,
                        "options": {"temperature": 0.3, "num_predict": 4096}
                    },
                    timeout=300
                )
                if r.status_code == 200:
                    return r.json().get("response", "").strip()
            except Exception as e:
                logger.warning(f"Tentativa {attempt+1}: {e}")
                time.sleep(5 * (attempt + 1))
        return None

# ============================================================================
# TRADUÇÃO E CORREÇÃO
# ============================================================================

def create_chunks(text: str, max_tokens: int = MAX_CHUNK_TOKENS) -> List[str]:
    """Divide texto em chunks."""
    if not text:
        return []

    chunks = []
    current = ""

    for para in text.split('\n\n'):
        para = para.strip()
        if not para:
            continue

        if count_tokens(current) + count_tokens(para) <= max_tokens:
            current += ('\n\n' if current else '') + para
        else:
            if current:
                chunks.append(current)
            current = para

    if current:
        chunks.append(current)

    return chunks


async def translate_text(client: OllamaClient, text: str, source_lang: str) -> str:
    """Traduz texto para português."""
    if source_lang == 'pt':
        return text

    chunks = create_chunks(text)
    translated = []

    lang_names = {'en': 'inglês', 'es': 'espanhol', 'fr': 'francês', 'de': 'alemão'}
    lang_name = lang_names.get(source_lang, source_lang)

    for i, chunk in enumerate(tqdm(chunks, desc="Traduzindo", unit="chunk")):
        # Verifica cache
        cached = cache.get(chunk)
        if cached:
            translated.append(cached)
            continue

        prompt = f"""Traduza o texto abaixo de {lang_name} para português brasileiro.
Mantenha o estilo literário. Retorne APENAS a tradução.

TEXTO:
{chunk}

TRADUÇÃO:"""

        result = await asyncio.get_event_loop().run_in_executor(
            None, client.generate, prompt
        )

        if result:
            cache.set(chunk, result, source_lang, client.model)
            translated.append(result)
        else:
            translated.append(chunk)  # Fallback

    return '\n\n'.join(translated)


async def correct_text(client: OllamaClient, text: str) -> str:
    """Corrige texto em português."""
    chunks = create_chunks(text)
    corrected = []

    for chunk in tqdm(chunks, desc="Corrigindo", unit="chunk"):
        prompt = f"""Corrija o texto abaixo. Corrija erros de gramática, ortografia e pontuação.
Mantenha o estilo original. Retorne APENAS o texto corrigido.

TEXTO:
{chunk}

TEXTO CORRIGIDO:"""

        result = await asyncio.get_event_loop().run_in_executor(
            None, client.generate, prompt
        )

        corrected.append(result if result else chunk)

    return '\n\n'.join(corrected)

# ============================================================================
# GERAÇÃO DOCX
# ============================================================================

def add_page_numbers(doc: Document):
    """Adiciona numeração de páginas."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()
        fld1 = OxmlElement('w:fldChar')
        fld1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.text = "PAGE"
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run._r.extend([fld1, instr, fld2])


def generate_docx(text: str, output_path: Path, title: str, author: str) -> bool:
    """Gera documento DOCX."""
    try:
        # Usa template se existir
        if TEMPLATE_DOCX.exists():
            doc = Document(str(TEMPLATE_DOCX))
            for p in list(doc.paragraphs):
                p._element.getparent().remove(p._element)
        else:
            doc = Document()

        # Título
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Autor
        if author:
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(author)
            author_run.italic = True
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Conteúdo
        chapter_re = re.compile(r'^(Capítulo|CAPÍTULO|Chapter|PARTE|LIVRO)\s+', re.I)

        for para_text in text.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            para = doc.add_paragraph()

            if chapter_re.match(para_text):
                run = para.add_run(para_text)
                run.bold = True
                run.font.size = Pt(18)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                para.add_run(para_text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_page_numbers(doc)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return True

    except Exception as e:
        logger.error(f"Erro ao gerar DOCX: {e}")
        return False

# ============================================================================
# PROCESSAMENTO PRINCIPAL
# ============================================================================

def extract_author_from_path(file_path: Path) -> Tuple[str, str]:
    """
    Extrai autor e título do caminho do arquivo.

    Padrões suportados:
    - CALIBRE: pasta/Titulo do Livro/arquivo.epub -> autor="Desconhecido", titulo="Titulo do Livro"
    - txt/Autor/livro.txt -> autor="Autor", titulo="livro"
    - Autor - Titulo.epub -> autor="Autor", titulo="Titulo"
    """
    title = file_path.stem
    author = "Desconhecido"

    # Padrão: "Autor - Titulo" no nome do arquivo
    if " - " in title:
        parts = title.split(" - ", 1)
        author = parts[0].strip()
        title = parts[1].strip()

    # Padrão: pasta pai é o autor (txt/Autor/livro.txt)
    elif file_path.parent.parent.name in ['txt', 'CALIBRE']:
        parent = file_path.parent.name
        # Se a pasta pai parece um título de livro, usa como título
        if file_path.parent.parent.name == 'CALIBRE':
            title = parent
        else:
            author = parent

    # Limpa caracteres problemáticos
    author = re.sub(r'[<>:"/\\|?*]', '', author)
    title = re.sub(r'[<>:"/\\|?*]', '', title)

    return author, title


async def process_book(file_path: Path, client: OllamaClient, output_dir: Path) -> bool:
    """Processa um livro completo."""
    # Extrai autor e título
    author, title = extract_author_from_path(file_path)
    book_id = f"{author}/{title}"

    logger.info(f"\n{'='*60}")
    logger.info(f"PROCESSANDO: {title}")
    logger.info(f"AUTOR: {author}")
    logger.info(f"{'='*60}")

    # 1. Extração
    result = extract_text_smart(file_path)

    if result.status != ExtractionStatus.SUCCESS:
        error_msg = result.error_message or result.status.value
        logger.error(f"✗ Extração falhou: {error_msg}")
        mark_as_failed(book_id, error_msg)
        return False

    text = result.text
    logger.info(f"✓ Extraído: {len(text)} caracteres")

    # 2. Detecção de idioma
    lang = detect_language(text)
    logger.info(f"✓ Idioma: {lang}")

    # 3. Tradução (se necessário)
    if lang != 'pt' and lang != 'unknown':
        logger.info(f"Traduzindo de {lang} para português...")
        text = await translate_text(client, text, lang)

    # 4. Correção
    logger.info("Corrigindo texto...")
    text = await correct_text(client, text)

    # 5. Geração DOCX em pasta do autor
    # Estrutura: docx/Autor/Titulo.docx
    author_dir = output_dir / author
    author_dir.mkdir(parents=True, exist_ok=True)
    output_path = author_dir / f"{title}.docx"

    if generate_docx(text, output_path, title, author):
        logger.info(f"✓ Salvo: {output_path}")
        mark_as_processed(book_id)
        return True
    else:
        mark_as_failed(book_id, "Erro ao gerar DOCX")
        return False

# ============================================================================
# MAIN
# ============================================================================

async def main_async():
    import argparse

    parser = argparse.ArgumentParser(description="Tradutor Inteligente de Livros")
    parser.add_argument('--input', '-i', type=Path, help='Arquivo para processar')
    parser.add_argument('--calibre', '-c', type=Path, default=CALIBRE_DIR)
    parser.add_argument('--txt', '-t', type=Path, default=INPUT_TXT_DIR, help='Pasta txt/')
    parser.add_argument('--output', '-o', type=Path, default=OUTPUT_DOCX_DIR)
    parser.add_argument('--model', '-m', default=OLLAMA_MODEL)
    parser.add_argument('--limit', '-l', type=int, default=0)

    args = parser.parse_args()

    logger.info("="*60)
    logger.info("TRADUTOR INTELIGENTE DE LIVROS")
    logger.info(f"Modelo: {args.model}")
    logger.info("="*60)

    # Verifica bibliotecas
    logger.info(f"PyMuPDF: {'✓' if HAS_PYMUPDF else '✗'}")
    logger.info(f"pdfplumber: {'✓' if HAS_PDFPLUMBER else '✗'}")
    logger.info(f"ebooklib: {'✓' if HAS_EBOOKLIB else '✗'}")

    try:
        client = OllamaClient(args.model)
    except Exception as e:
        logger.error(f"Falha ao conectar: {e}")
        return 1

    processed = load_processed_books()
    failed = load_failed_books()
    logger.info(f"Já processados: {len(processed)} | Falhos: {len(failed)}")

    # Arquivo único
    if args.input:
        success = await process_book(args.input, client, args.output)
        return 0 if success else 1

    # Encontra livros (CALIBRE e txt/)
    books = []

    # Da pasta txt/
    if args.txt.exists():
        for author_dir in args.txt.iterdir():
            if author_dir.is_dir():
                for f in author_dir.glob('*.txt'):
                    book_id = f"{author_dir.name}/{f.stem}"
                    if book_id not in processed:
                        books.append((book_id, f))

    # Da pasta CALIBRE
    if args.calibre.exists():
        for folder in args.calibre.iterdir():
            if folder.is_dir():
                book_id = folder.name
                if book_id in processed:
                    continue

                for ext in ['.epub', '.pdf', '.html', '.txt']:
                    files = list(folder.glob(f'*{ext}'))
                    if files:
                        books.append((book_id, files[0]))
                        break

    logger.info(f"Livros para processar: {len(books)}")

    if args.limit > 0:
        books = books[:args.limit]

    success_count = 0
    fail_count = 0

    for book_id, book_path in books:
        try:
            if await process_book(book_path, client, args.output):
                success_count += 1
            else:
                fail_count += 1
        except Exception as e:
            logger.error(f"Erro em {book_id}: {e}")
            mark_as_failed(book_id, str(e))
            fail_count += 1

    logger.info("\n" + "="*60)
    logger.info(f"CONCLUÍDO - Sucesso: {success_count} | Falhas: {fail_count}")
    logger.info("="*60)

    return 0


def main():
    return asyncio.run(main_async())


if __name__ == "__main__":
    sys.exit(main())
