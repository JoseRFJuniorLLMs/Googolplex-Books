# -*- coding: utf-8 -*-
"""
PROCESSOR.PY - Processador de Livros com Modelos Locais
========================================================
Versão melhorada do correcaoKDP.py que usa Ollama (modelos locais)
como backend principal, com fallback para Gemini/OpenAI.

Recursos:
- Suporte a Ollama (qwen2.5, llama3.1, etc.)
- Cache de correções em SQLite
- Processamento paralelo
- Geração de DOCX formatado
- Identificação de notas de rodapé
- Integração com base de dados (marca como processado)
"""

import os
import re
import sys
import time
import asyncio
import hashlib
import sqlite3
import logging
import traceback
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Tuple, Dict
from concurrent.futures import ThreadPoolExecutor

import requests
from tqdm import tqdm

# python-docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configurações
sys.path.insert(0, str(Path(__file__).parent.parent))
from config.settings import (
    DATABASE_PATH, CACHE_PATH, LOG_DIR, TEMPLATE_DOCX,
    INPUT_TXT_DIR, TRANSLATED_DIR, OUTPUT_DOCX_DIR,
    MODEL_BACKEND, OLLAMA_BASE_URL, OLLAMA_MODEL,
    GEMINI_API_KEY, OPENAI_API_KEY,
    MAX_CHUNK_TOKENS, MAX_OUTPUT_TOKENS, TEMPERATURE,
    MAX_RETRIES, RETRY_BASE_WAIT, PARALLEL_CHUNKS,
    CHAPTER_PATTERNS, PAGE_BREAK_MARKER, AI_FAILURE_MARKER
)

from database import BooksDatabase

# ============================================================================
# LOGGING
# ============================================================================

LOG_DIR.mkdir(exist_ok=True)
log_file = LOG_DIR / f"processor_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)-8s | %(funcName)s | %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# CACHE DE CORREÇÕES
# ============================================================================

class CorrectionCache:
    """Cache SQLite para correções já feitas."""

    def __init__(self, db_path: Path = CACHE_PATH):
        self.db_path = db_path
        self.db_path.parent.mkdir(exist_ok=True)
        self._init_db()

    def _init_db(self):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS corrections (
                    hash TEXT PRIMARY KEY,
                    original TEXT,
                    corrected TEXT,
                    model TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('''
                CREATE TABLE IF NOT EXISTS footnotes (
                    hash TEXT PRIMARY KEY,
                    marked_text TEXT,
                    model TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_corrections ON corrections(hash)')

    def _hash(self, text: str) -> str:
        return hashlib.sha256(text.encode()).hexdigest()[:32]

    def get_correction(self, text: str) -> Optional[str]:
        with sqlite3.connect(self.db_path) as conn:
            cur = conn.execute(
                'SELECT corrected FROM corrections WHERE hash = ?',
                (self._hash(text),)
            )
            row = cur.fetchone()
            return row[0] if row else None

    def set_correction(self, original: str, corrected: str, model: str):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                INSERT OR REPLACE INTO corrections (hash, original, corrected, model)
                VALUES (?, ?, ?, ?)
            ''', (self._hash(original), original, corrected, model))

    def get_footnotes(self, text: str) -> Optional[str]:
        with sqlite3.connect(self.db_path) as conn:
            cur = conn.execute(
                'SELECT marked_text FROM footnotes WHERE hash = ?',
                (self._hash(text),)
            )
            row = cur.fetchone()
            return row[0] if row else None

    def set_footnotes(self, original: str, marked: str, model: str):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                INSERT OR REPLACE INTO footnotes (hash, marked_text, model)
                VALUES (?, ?, ?)
            ''', (self._hash(original), marked, model))

cache = CorrectionCache()

# ============================================================================
# CONTAGEM DE TOKENS
# ============================================================================

try:
    import tiktoken
    _encoder = tiktoken.get_encoding("cl100k_base")

    def count_tokens(text: str) -> int:
        return len(_encoder.encode(text)) if text else 0

    logger.info("Usando tiktoken para contagem de tokens")
except ImportError:
    def count_tokens(text: str) -> int:
        return len(text) // 4 if text else 0

    logger.warning("tiktoken não instalado - usando estimativa")

# ============================================================================
# CLIENTE DE MODELOS
# ============================================================================

class ModelClient:
    """Cliente unificado para diferentes backends de IA."""

    def __init__(self, backend: str = MODEL_BACKEND):
        self.backend = backend
        self.model_name = OLLAMA_MODEL
        self._validate()

    def _validate(self):
        """Valida conexão com o backend."""
        if self.backend == "ollama":
            try:
                r = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=5)
                if r.status_code == 200:
                    models = [m['name'] for m in r.json().get('models', [])]
                    base_model = OLLAMA_MODEL.split(':')[0]
                    if not any(base_model in m for m in models):
                        logger.warning(f"Modelo {OLLAMA_MODEL} não encontrado")
                        logger.info(f"Modelos disponíveis: {models}")
                        logger.info(f"Execute: ollama pull {OLLAMA_MODEL}")
                    else:
                        logger.info(f"Ollama OK - Modelo: {OLLAMA_MODEL}")
                else:
                    raise ConnectionError()
            except Exception as e:
                logger.error(f"Ollama não disponível em {OLLAMA_BASE_URL}")
                logger.info("Instale: https://ollama.com")
                raise

        elif self.backend == "gemini":
            if not GEMINI_API_KEY:
                raise ValueError("GOOGLE_API_KEY não configurada")
            import google.generativeai as genai
            genai.configure(api_key=GEMINI_API_KEY)
            self._gemini = genai.GenerativeModel("gemini-1.5-pro")
            self.model_name = "gemini-1.5-pro"
            logger.info("Usando Gemini API")

        elif self.backend == "openai":
            if not OPENAI_API_KEY:
                raise ValueError("OPENAI_API_KEY não configurada")
            from openai import OpenAI
            self._openai = OpenAI(api_key=OPENAI_API_KEY)
            self.model_name = "gpt-4o-mini"
            logger.info("Usando OpenAI API")

    def generate(self, prompt: str, max_retries: int = MAX_RETRIES) -> Optional[str]:
        """Gera texto com retry."""
        for attempt in range(max_retries):
            try:
                if self.backend == "ollama":
                    return self._gen_ollama(prompt)
                elif self.backend == "gemini":
                    return self._gen_gemini(prompt)
                elif self.backend == "openai":
                    return self._gen_openai(prompt)
            except Exception as e:
                wait = RETRY_BASE_WAIT * (2 ** attempt)
                logger.warning(f"Tentativa {attempt+1}/{max_retries} falhou: {e}")
                if attempt < max_retries - 1:
                    time.sleep(wait)

        logger.error("Falha após todas as tentativas")
        return None

    def _gen_ollama(self, prompt: str) -> str:
        r = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": TEMPERATURE,
                    "num_predict": MAX_OUTPUT_TOKENS
                }
            },
            timeout=300
        )
        r.raise_for_status()
        return r.json().get("response", "").strip()

    def _gen_gemini(self, prompt: str) -> str:
        response = self._gemini.generate_content(prompt)
        if response.candidates:
            parts = response.candidates[0].content.parts
            return "".join(p.text for p in parts if hasattr(p, 'text')).strip()
        return ""

    def _gen_openai(self, prompt: str) -> str:
        r = self._openai.chat.completions.create(
            model=self.model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=TEMPERATURE,
            max_tokens=MAX_OUTPUT_TOKENS
        )
        return r.choices[0].message.content.strip()

# ============================================================================
# CHUNKING
# ============================================================================

def create_chunks(text: str, max_tokens: int = MAX_CHUNK_TOKENS) -> List[str]:
    """Divide texto em chunks respeitando limites."""
    if not text or not text.strip():
        return []

    chunks = []
    current = ""
    current_tokens = 0

    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue

        para_tokens = count_tokens(para)

        if current_tokens + para_tokens + 2 <= max_tokens:
            current += ("\n\n" if current else "") + para
            current_tokens = count_tokens(current)
        else:
            if current:
                chunks.append(current)

            if para_tokens > max_tokens:
                # Subdivide parágrafos grandes
                sentences = re.split(r'(?<=[.!?])\s+', para)
                sub = ""
                sub_tokens = 0

                for sent in sentences:
                    sent_tokens = count_tokens(sent)
                    if sub_tokens + sent_tokens + 1 <= max_tokens:
                        sub += (" " if sub else "") + sent
                        sub_tokens = count_tokens(sub)
                    else:
                        if sub:
                            chunks.append(sub)
                        if sent_tokens > max_tokens:
                            chunks.append(sent)
                            sub = ""
                            sub_tokens = 0
                        else:
                            sub = sent
                            sub_tokens = sent_tokens

                current = sub if sub else ""
                current_tokens = sub_tokens if sub else 0
            else:
                current = para
                current_tokens = para_tokens

    if current:
        chunks.append(current)

    # Merge chunks pequenos
    merged = []
    temp = ""
    temp_tokens = 0

    for chunk in chunks:
        chunk_tokens = count_tokens(chunk)
        if temp_tokens + chunk_tokens + 2 <= max_tokens:
            temp += ("\n\n" if temp else "") + chunk
            temp_tokens = count_tokens(temp)
        else:
            if temp:
                merged.append(temp)
            temp = chunk
            temp_tokens = chunk_tokens

    if temp:
        merged.append(temp)

    logger.info(f"Texto dividido em {len(merged)} chunks")
    return merged

# ============================================================================
# PROMPTS
# ============================================================================

def get_correction_prompt(chunk: str, author: str, is_first: bool = False) -> str:
    """Prompt para correção de OCR."""
    ctx = "Início do livro." if is_first else "Continuação do texto."

    return f"""Você é um editor literário profissional. {ctx}

TAREFA: Corrija o texto abaixo (OCR de livro do autor {author}).

REGRAS:
1. Corrija TODOS os erros: ortografia, gramática, pontuação, acentuação
2. Corrija erros de OCR: 'rn'→'m', 'cl'→'d', 'l'→'i', espaços errados
3. MANTENHA estrutura de parágrafos (separados por linha em branco)
4. MANTENHA estilo e tom do autor
5. NÃO adicione/remova conteúdo
6. NÃO use markdown
7. Retorne APENAS o texto corrigido

TEXTO:
\"\"\"
{chunk}
\"\"\"

CORRIGIDO:"""


def get_footnote_prompt(chunk: str, author: str) -> str:
    """Prompt para identificar notas de rodapé."""
    return f"""Analise o texto do autor {author} e insira marcadores de nota APENAS onde necessário.

INSERIR MARCADORES PARA:
- Termos em idioma estrangeiro (latim, francês, etc.)
- Nomes próprios pouco conhecidos
- Termos técnicos que precisam de explicação
- Citações/referências

FORMATO: [NOTA:termo|explicação breve]

EXEMPLO:
"o conceito de Übermensch[NOTA:Übermensch|Super-homem, conceito de Nietzsche]"

REGRAS:
- NÃO altere o texto
- Seja conservador - só termos realmente necessários
- Explicações breves (máx 20 palavras)
- Mantenha parágrafos

TEXTO:
\"\"\"
{chunk}
\"\"\"

COM MARCADORES:"""

# ============================================================================
# PROCESSAMENTO
# ============================================================================

async def process_correction(client: ModelClient, chunk: str, author: str,
                            is_first: bool, idx: int) -> Tuple[int, str]:
    """Processa chunk para correção."""
    # Verifica cache
    cached = cache.get_correction(chunk)
    if cached:
        logger.debug(f"Chunk {idx}: cache hit")
        return (idx, cached)

    prompt = get_correction_prompt(chunk, author, is_first)
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, client.generate, prompt)

    if result:
        cache.set_correction(chunk, result, client.model_name)
        return (idx, result)
    else:
        return (idx, f"{AI_FAILURE_MARKER}\n\n{chunk}")


async def process_footnotes(client: ModelClient, chunk: str, author: str,
                           idx: int) -> Tuple[int, str]:
    """Processa chunk para notas de rodapé."""
    cached = cache.get_footnotes(chunk)
    if cached:
        return (idx, cached)

    prompt = get_footnote_prompt(chunk, author)
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, client.generate, prompt)

    if result:
        cache.set_footnotes(chunk, result, client.model_name)
        return (idx, result)
    return (idx, chunk)


async def process_chunks_parallel(client: ModelClient, chunks: List[str],
                                 author: str, mode: str = "correction") -> List[str]:
    """Processa chunks em paralelo."""
    results = [None] * len(chunks)
    batch_size = PARALLEL_CHUNKS

    desc = "Corrigindo" if mode == "correction" else "Notas"

    with tqdm(total=len(chunks), desc=desc, unit="chunk") as pbar:
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i+batch_size]
            tasks = []

            for j, chunk in enumerate(batch):
                idx = i + j
                if mode == "correction":
                    tasks.append(process_correction(client, chunk, author, idx == 0, idx))
                else:
                    tasks.append(process_footnotes(client, chunk, author, idx))

            batch_results = await asyncio.gather(*tasks, return_exceptions=True)

            for result in batch_results:
                if isinstance(result, Exception):
                    logger.error(f"Erro: {result}")
                else:
                    idx, text = result
                    results[idx] = text
                pbar.update(1)

    # Fallback para None
    for i, r in enumerate(results):
        if r is None:
            results[i] = chunks[i]

    return results

# ============================================================================
# DOCX
# ============================================================================

def extract_footnotes(text: str) -> Tuple[str, List[Dict]]:
    """Extrai marcadores [NOTA:...|...] do texto."""
    pattern = r'\[NOTA:([^|]+)\|([^\]]+)\]'
    footnotes = []
    counter = 1

    def replace(m):
        nonlocal counter
        footnotes.append({
            'number': counter,
            'term': m.group(1).strip(),
            'explanation': m.group(2).strip()
        })
        result = f"[{counter}]"
        counter += 1
        return result

    clean = re.sub(pattern, replace, text)
    return clean, footnotes


def add_page_numbers(doc: Document):
    """Adiciona numeração de páginas."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        if not footer.paragraphs:
            para = footer.add_paragraph()
        else:
            para = footer.paragraphs[0]

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()

        # Campo PAGE
        for el in ['begin', None, 'separate', 'end']:
            if el:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), el)
                run._r.append(fc)
            else:
                it = OxmlElement('w:instrText')
                it.set(qn('xml:space'), 'preserve')
                it.text = "PAGE"
                run._r.append(it)


def generate_docx(text: str, footnotes: List[Dict], template_path: Path,
                  output_path: Path, author: str, book_name: str) -> bool:
    """Gera documento DOCX formatado."""
    try:
        if not template_path.exists():
            logger.error(f"Template não encontrado: {template_path}")
            return False

        doc = Document(template_path)

        # Limpa conteúdo
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)

        chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)

        # Processa texto
        parts = text.split(PAGE_BREAK_MARKER)

        for part_idx, part in enumerate(parts):
            if part_idx > 0:
                doc.add_page_break()

            for para_text in part.strip().split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue

                is_chapter = chapter_regex.match(para_text) is not None
                is_error = para_text.startswith(AI_FAILURE_MARKER)

                para = doc.add_paragraph()

                # Processa marcadores [N]
                note_pattern = r'\[(\d+)\]'
                last_end = 0

                for match in re.finditer(note_pattern, para_text):
                    if match.start() > last_end:
                        run = para.add_run(para_text[last_end:match.start()])
                        if is_chapter:
                            run.font.size = Pt(24)
                            run.bold = True
                        elif is_error:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.italic = True

                    # Nota sobrescrita
                    note_run = para.add_run(match.group(0))
                    note_run.font.superscript = True
                    note_run.font.size = Pt(8)

                    last_end = match.end()

                if last_end < len(para_text):
                    run = para.add_run(para_text[last_end:])
                    if is_chapter:
                        run.font.size = Pt(24)
                        run.bold = True
                    elif is_error:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.italic = True

                if is_chapter:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.space_before = Pt(24)
                    para.space_after = Pt(12)
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Seção de notas
        if footnotes:
            doc.add_page_break()
            title = doc.add_paragraph("NOTAS")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(14)

            doc.add_paragraph()

            for note in footnotes:
                p = doc.add_paragraph()
                p.add_run(f"[{note['number']}] ").bold = True
                p.add_run(f"{note['term']}: ").italic = True
                p.add_run(note['explanation'])

        # Numeração de páginas
        add_page_numbers(doc)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"DOCX salvo: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Erro ao gerar DOCX: {e}")
        logger.error(traceback.format_exc())
        return False

# ============================================================================
# MAIN
# ============================================================================

async def process_book(input_path: Path, output_path: Path, template: Path,
                       author: str, book_name: str, client: ModelClient,
                       book_id: str = None) -> bool:
    """Processa um livro completo."""
    logger.info("="*60)
    logger.info(f"Processando: {author} - {book_name}")
    logger.info(f"Entrada: {input_path}")
    logger.info(f"Saída: {output_path}")
    logger.info("="*60)

    start = time.time()

    try:
        # 1. Lê arquivo
        with open(input_path, 'r', encoding='utf-8') as f:
            original = f.read()

        logger.info(f"Texto: {len(original):,} chars, ~{count_tokens(original):,} tokens")

        # 2. Chunks
        chunks = create_chunks(original)
        if not chunks:
            logger.error("Nenhum chunk gerado")
            return False

        # 3. Correção
        logger.info(f"Corrigindo {len(chunks)} chunks...")
        corrected = await process_chunks_parallel(client, chunks, author, "correction")
        corrected_text = "\n\n".join(corrected)

        # 4. Notas de rodapé
        logger.info("Identificando notas...")
        note_chunks = create_chunks(corrected_text)
        marked = await process_chunks_parallel(client, note_chunks, author, "footnotes")
        marked_text = "\n\n".join(marked)

        # 5. Extrai notas
        final_text, footnotes = extract_footnotes(marked_text)
        logger.info(f"Notas encontradas: {len(footnotes)}")

        # 6. Gera DOCX
        success = generate_docx(final_text, footnotes, template, output_path, author, book_name)

        elapsed = time.time() - start

        if success:
            logger.info(f"Concluído em {elapsed:.1f}s")

            # Atualiza base de dados se tiver book_id
            if book_id:
                try:
                    with BooksDatabase() as db:
                        db.mark_as_processed(book_id, str(output_path))
                except Exception as e:
                    logger.warning(f"Erro ao atualizar DB: {e}")

        return success

    except Exception as e:
        logger.error(f"Erro: {e}")
        logger.error(traceback.format_exc())
        return False


async def main_async():
    """Main assíncrona."""
    import argparse

    parser = argparse.ArgumentParser(description="Processador de Livros - Modelos Locais")
    parser.add_argument('--input', '-i', type=Path, help='Arquivo TXT de entrada')
    parser.add_argument('--output', '-o', type=Path, help='Arquivo DOCX de saída')
    parser.add_argument('--author', '-a', default='Desconhecido', help='Nome do autor')
    parser.add_argument('--backend', '-b', choices=['ollama', 'gemini', 'openai'],
                       default=MODEL_BACKEND, help='Backend de IA')
    parser.add_argument('--model', '-m', help='Modelo específico (ex: qwen2.5:14b)')
    parser.add_argument('--batch', action='store_true', help='Processa todos em txt/')
    parser.add_argument('--book-id', help='ID do livro no banco')

    args = parser.parse_args()

    # Ajusta modelo se especificado
    global OLLAMA_MODEL
    if args.model:
        OLLAMA_MODEL = args.model

    logger.info("="*60)
    logger.info("PROCESSOR - Modelos Locais")
    logger.info(f"Backend: {args.backend}")
    logger.info(f"Modelo: {OLLAMA_MODEL if args.backend == 'ollama' else args.backend}")
    logger.info("="*60)

    try:
        client = ModelClient(args.backend)
    except Exception as e:
        logger.error(f"Falha ao iniciar cliente: {e}")
        return 1

    # Modo arquivo único
    if args.input:
        if not args.input.exists():
            logger.error(f"Arquivo não encontrado: {args.input}")
            return 1

        output = args.output or args.input.with_suffix('.docx')
        book_name = args.input.stem

        success = await process_book(
            args.input, output, TEMPLATE_DOCX,
            args.author, book_name, client, args.book_id
        )
        return 0 if success else 1

    # Modo batch - processa livros traduzidos (pasta translated/)
    if args.batch:
        logger.info(f"Modo batch: {TRANSLATED_DIR}")

        if not TRANSLATED_DIR.exists():
            logger.error(f"Diretório não existe: {TRANSLATED_DIR}")
            return 1

        # Busca todos os arquivos *_pt.txt na pasta translated
        pt_files = list(TRANSLATED_DIR.rglob("*_pt.txt"))

        if not pt_files:
            logger.warning("Nenhum arquivo traduzido encontrado (*_pt.txt)")
            return 0

        logger.info(f"Encontrados {len(pt_files)} arquivos traduzidos")

        success_count = 0
        fail_count = 0

        for txt_file in sorted(pt_files):
            # Extrai nome do livro (remove _pt do final)
            book_name = txt_file.stem.replace('_pt', '')

            # Usa pasta pai como autor (ou "Desconhecido" se estiver na raiz)
            author_name = txt_file.parent.name if txt_file.parent != TRANSLATED_DIR else "Desconhecido"

            # Define saída
            output_dir = OUTPUT_DOCX_DIR / author_name
            output_file = output_dir / f"{book_name}_Final.docx"

            # Pula se já existe
            if output_file.exists():
                logger.info(f"Já existe: {output_file.name}")
                continue

            logger.info(f"\n--- Processando: {book_name} ({author_name}) ---")

            ok = await process_book(
                txt_file, output_file, TEMPLATE_DOCX,
                author_name, book_name, client
            )

            if ok:
                success_count += 1
            else:
                fail_count += 1

        logger.info("\n" + "="*60)
        logger.info(f"RESUMO: Sucesso={success_count}, Falhas={fail_count}")
        logger.info("="*60)

        return 0 if fail_count == 0 else 1

    parser.print_help()
    return 0


def main():
    """Ponto de entrada."""
    try:
        return asyncio.run(main_async())
    except KeyboardInterrupt:
        logger.warning("\nInterrompido")
        return 130


if __name__ == "__main__":
    sys.exit(main())
