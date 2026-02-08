# -*- coding: utf-8 -*-
"""
PROCESSOR_BILINGUAL.PY - Pipeline 2: Versão Bilíngue + Semantic Priming
=========================================================================
Gera versão bilíngue com 100 palavras-chave em inglês + exemplos.

Processo:
1. Análise TF-IDF (100 palavras mais importantes)
2. Clustering semântico (agrupa palavras relacionadas)
3. Substituição PT → EN (100 palavras)
4. Geração de 3 frases exemplo em inglês
5. DOCX com texto bilíngue + notas de rodapé

Entrada: translated/[Autor]/[Titulo]_pt.txt
Saída: docx/pipeline2/[Autor]/[Titulo]_pt_bilingual.docx
"""

import os
import re
import sys
import time
import logging
import traceback
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional
from collections import Counter
import json

# Configurações
sys.path.insert(0, str(Path(__file__).parent.parent))
from config.settings import (
    OUTPUT_DOCX_DIR, TEMPLATE_DOCX, OLLAMA_BASE_URL, OLLAMA_MODEL,
    TEMPERATURE, MAX_OUTPUT_TOKENS, LOG_DIR
)

# ============================================================================
# LOGGING
# ============================================================================

LOG_DIR.mkdir(exist_ok=True)
log_file = LOG_DIR / f"bilingual_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)-8s | %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# STOP WORDS (PT)
# ============================================================================

STOP_WORDS_PT = set([
    'o', 'a', 'os', 'as', 'um', 'uma', 'uns', 'umas',
    'de', 'do', 'da', 'dos', 'das',
    'em', 'no', 'na', 'nos', 'nas',
    'por', 'para', 'com', 'sem',
    'e', 'ou', 'mas', 'porém', 'contudo',
    'que', 'qual', 'quais', 'quando', 'onde',
    'se', 'é', 'era', 'foi', 'ser', 'estar',
    'ter', 'há', 'ao', 'aos', 'à', 'às',
    'ele', 'ela', 'eles', 'elas', 'seu', 'sua',
    'este', 'esta', 'esse', 'essa', 'aquele', 'aquela',
    'mais', 'menos', 'muito', 'muitos', 'pouco', 'poucos',
    'todo', 'toda', 'todos', 'todas', 'cada',
    'já', 'ainda', 'também', 'nem', 'não', 'sim',
])

# ============================================================================
# ANÁLISE TF-IDF SIMPLIFICADA
# ============================================================================

def extract_keywords_tfidf(text: str, num_keywords: int = 100) -> List[Tuple[str, float]]:
    """
    Extrai palavras-chave usando TF-IDF simplificado.

    Returns:
        Lista de (palavra, score) ordenada por relevância
    """
    # Tokenização simples
    words = re.findall(r'\b[a-záàâãéêíóôõúçA-ZÁÀÂÃÉÊÍÓÔÕÚÇ]+\b', text.lower())

    # Remove stop words e palavras muito curtas
    words_filtered = [
        w for w in words
        if w not in STOP_WORDS_PT and len(w) > 3
    ]

    # Contagem de frequência (TF)
    word_freq = Counter(words_filtered)

    # Score simples: frequência * comprimento da palavra
    # (palavras mais longas tendem a ser mais específicas)
    word_scores = {
        word: freq * (len(word) ** 0.5)
        for word, freq in word_freq.items()
    }

    # Ordena por score e retorna top N
    sorted_words = sorted(word_scores.items(), key=lambda x: x[1], reverse=True)

    return sorted_words[:num_keywords]

# ============================================================================
# TRADUÇÃO PT → EN
# ============================================================================

# Dicionário básico PT → EN (expandível)
PT_TO_EN = {
    # Verbos comuns
    'ser': 'be', 'estar': 'be', 'ter': 'have', 'fazer': 'do', 'poder': 'can',
    'dizer': 'say', 'ver': 'see', 'saber': 'know', 'querer': 'want',
    'dar': 'give', 'ir': 'go', 'vir': 'come', 'falar': 'speak',

    # Substantivos comuns
    'homem': 'man', 'mulher': 'woman', 'pessoa': 'person', 'gente': 'people',
    'tempo': 'time', 'dia': 'day', 'ano': 'year', 'vida': 'life',
    'mundo': 'world', 'casa': 'house', 'lugar': 'place', 'coisa': 'thing',
    'modo': 'way', 'trabalho': 'work', 'mão': 'hand', 'olho': 'eye',
    'cabeça': 'head', 'parte': 'part', 'nome': 'name', 'número': 'number',

    # Adjetivos comuns
    'grande': 'big', 'pequeno': 'small', 'bom': 'good', 'mau': 'bad',
    'novo': 'new', 'velho': 'old', 'outro': 'other', 'mesmo': 'same',
    'próprio': 'own', 'certo': 'right', 'último': 'last', 'longo': 'long',
}

def translate_word_pt_to_en(word_pt: str, use_ai: bool = False) -> str:
    """
    Traduz palavra PT → EN.

    Args:
        word_pt: Palavra em português
        use_ai: Se True, usa IA para traduzir (mais lento)

    Returns:
        Palavra em inglês
    """
    word_lower = word_pt.lower()

    # Tenta dicionário básico primeiro
    if word_lower in PT_TO_EN:
        return PT_TO_EN[word_lower]

    # TODO: Implementar tradução via IA (Ollama)
    # Por enquanto, retorna transliteração simples
    return word_pt

# ============================================================================
# GERAÇÃO DE EXEMPLOS
# ============================================================================

def generate_example_sentences(word_en: str, word_pt: str, model_url: str, model_name: str) -> List[str]:
    """
    Gera 3 frases exemplo em inglês usando a palavra.

    Args:
        word_en: Palavra em inglês
        word_pt: Palavra em português (para contexto)
        model_url: URL do Ollama
        model_name: Nome do modelo

    Returns:
        Lista com 3 frases exemplo
    """
    try:
        import requests

        prompt = f"""Generate 3 simple English example sentences using the word "{word_en}".

RULES:
- Each sentence must be 8-15 words long
- Use simple, clear language
- Show different contexts
- Do NOT add numbers or labels
- Return only the 3 sentences, one per line

Examples for "king":
The king ruled for forty years.
He became king at age twenty one.
A wise king makes good decisions.

Now generate for "{word_en}":"""

        response = requests.post(
            f"{model_url}/api/generate",
            json={
                "model": model_name,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.3,
                    "num_predict": 200,
                }
            },
            timeout=60
        )

        if response.status_code == 200:
            result = response.json().get('response', '').strip()

            # Extrai frases (uma por linha)
            sentences = [
                line.strip()
                for line in result.split('\n')
                if line.strip() and len(line.strip()) > 10
            ]

            # Limpa pontos, números, etc
            sentences = [
                re.sub(r'^\d+[\.\)]\s*', '', s).strip()
                for s in sentences
            ]

            # Pega apenas 3
            sentences = sentences[:3]

            # Se não gerou 3, cria genéricas
            while len(sentences) < 3:
                sentences.append(f"This is an example with {word_en}.")

            return sentences[:3]

    except Exception as e:
        logger.warning(f"Erro ao gerar exemplos para '{word_en}': {e}")

    # Fallback: exemplos genéricos
    return [
        f"This is an example with {word_en}.",
        f"The word {word_en} is important.",
        f"We use {word_en} in this context.",
    ]

# ============================================================================
# SUBSTITUIÇÃO NO TEXTO
# ============================================================================

def substitute_keywords(text: str, keywords: List[str], translations: Dict[str, str]) -> Tuple[str, List[Dict]]:
    """
    Substitui palavras PT → EN no texto e cria marcadores [N].

    Args:
        text: Texto em português
        keywords: Lista de palavras para substituir
        translations: Dicionário PT → EN

    Returns:
        (texto_marcado, lista_de_notas)
    """
    marked_text = text
    footnotes = []
    note_counter = 1

    # Ordena palavras por tamanho (maior primeiro) para evitar substituições parciais
    keywords_sorted = sorted(keywords, key=len, reverse=True)

    for word_pt in keywords_sorted:
        word_en = translations.get(word_pt, word_pt)

        # Padrão: palavra completa (não substituir partes de palavras)
        pattern = r'\b' + re.escape(word_pt) + r'\b'

        # Verifica se palavra existe no texto
        if re.search(pattern, marked_text, re.IGNORECASE):
            # Substitui primeira ocorrência e adiciona marcador
            def replacer(match):
                nonlocal note_counter
                footnotes.append({
                    'number': note_counter,
                    'word_pt': word_pt,
                    'word_en': word_en,
                    'examples': []  # Será preenchido depois
                })
                replacement = f"{word_en}[{note_counter}]"
                note_counter += 1
                return replacement

            # Substitui apenas primeira ocorrência
            marked_text = re.sub(pattern, replacer, marked_text, count=1, flags=re.IGNORECASE)

    return marked_text, footnotes

# ============================================================================
# GERAÇÃO DOCX
# ============================================================================

def generate_bilingual_docx(
    text: str,
    footnotes: List[Dict],
    template_path: Path,
    output_path: Path,
    author: str,
    book_name: str
) -> bool:
    """Gera DOCX bilíngue com notas de rodapé."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not template_path.exists():
            logger.error(f"Template não encontrado: {template_path}")
            return False

        doc = Document(template_path)

        # Limpa conteúdo
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)

        # Adiciona texto bilíngue
        for para_text in text.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            para = doc.add_paragraph()

            # Processa marcadores [N]
            note_pattern = r'\[(\d+)\]'
            last_end = 0

            for match in re.finditer(note_pattern, para_text):
                if match.start() > last_end:
                    run = para.add_run(para_text[last_end:match.start()])

                # Nota sobrescrita
                note_run = para.add_run(match.group(0))
                note_run.font.superscript = True
                note_run.font.size = Pt(8)

                last_end = match.end()

            if last_end < len(para_text):
                para.add_run(para_text[last_end:])

            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Seção de notas
        if footnotes:
            doc.add_page_break()
            title = doc.add_paragraph("NOTAS - VOCABULARY")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(14)

            doc.add_paragraph()

            for note in footnotes:
                p = doc.add_paragraph()

                # [N] word_en (PT: word_pt)
                p.add_run(f"[{note['number']}] ").bold = True
                p.add_run(f"{note['word_en']} ").font.size = Pt(12)
                p.add_run(f"(PT: {note['word_pt']})").italic = True

                # 3 exemplos
                if note.get('examples'):
                    for example in note['examples']:
                        ex_para = doc.add_paragraph(f"  • {example}", style='List Bullet')
                        ex_para.paragraph_format.left_indent = Inches(0.5)

                # Espaço entre notas
                doc.add_paragraph()

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"DOCX salvo: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Erro ao gerar DOCX: {e}")
        logger.error(traceback.format_exc())
        return False

# ============================================================================
# PROCESSAMENTO PRINCIPAL
# ============================================================================

def process_bilingual(
    input_path: Path,
    output_path: Path,
    author: str,
    book_name: str,
    num_keywords: int = 100
) -> bool:
    """
    Processa arquivo para versão bilíngue.

    Args:
        input_path: Arquivo PT traduzido
        output_path: Arquivo DOCX bilíngue
        author: Nome do autor
        book_name: Nome do livro
        num_keywords: Número de palavras-chave (padrão: 100)

    Returns:
        True se sucesso
    """
    logger.info("="*70)
    logger.info(f"Processando: {author} - {book_name}")
    logger.info(f"Entrada: {input_path}")
    logger.info(f"Saída: {output_path}")
    logger.info("="*70)

    start = time.time()

    try:
        # 1. Lê arquivo
        logger.info("1/6 - Lendo arquivo...")
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()

        logger.info(f"  Texto: {len(text):,} caracteres")

        # 2. Extrai palavras-chave (TF-IDF)
        logger.info(f"2/6 - Extraindo {num_keywords} palavras-chave (TF-IDF)...")
        keywords_with_scores = extract_keywords_tfidf(text, num_keywords)
        keywords = [word for word, score in keywords_with_scores]

        logger.info(f"  Top 10: {', '.join(keywords[:10])}")

        # 3. Traduz palavras PT → EN
        logger.info("3/6 - Traduzindo palavras PT → EN...")
        translations = {}
        for word_pt in keywords:
            word_en = translate_word_pt_to_en(word_pt, use_ai=False)
            translations[word_pt] = word_en

        logger.info(f"  Traduções: {len(translations)}")

        # 4. Substitui no texto
        logger.info("4/6 - Substituindo palavras no texto...")
        marked_text, footnotes = substitute_keywords(text, keywords, translations)

        logger.info(f"  Substituições: {len(footnotes)}")

        # 5. Gera exemplos para cada palavra
        logger.info("5/6 - Gerando exemplos (3 por palavra)...")
        logger.info("  (Isso vai demorar... usando IA para cada palavra)")

        for i, note in enumerate(footnotes, 1):
            logger.info(f"  {i}/{len(footnotes)}: {note['word_en']}...")
            examples = generate_example_sentences(
                note['word_en'],
                note['word_pt'],
                OLLAMA_BASE_URL,
                OLLAMA_MODEL
            )
            note['examples'] = examples

        logger.info("  Exemplos gerados!")

        # 6. Gera DOCX
        logger.info("6/6 - Gerando DOCX bilíngue...")
        success = generate_bilingual_docx(
            marked_text,
            footnotes,
            TEMPLATE_DOCX,
            output_path,
            author,
            book_name
        )

        elapsed = time.time() - start

        if success:
            logger.info(f"Concluído em {elapsed/60:.1f} minutos")
            logger.info(f"  Palavras substituídas: {len(footnotes)}")
            logger.info(f"  Total de exemplos: {len(footnotes) * 3}")

        return success

    except Exception as e:
        logger.error(f"Erro: {e}")
        logger.error(traceback.format_exc())
        return False

# ============================================================================
# MAIN
# ============================================================================

def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="Pipeline 2 - Processador Bilíngue + Semantic Priming"
    )
    parser.add_argument('--input', '-i', type=Path,
                       help='Arquivo TXT traduzido (PT)')
    parser.add_argument('--output', '-o', type=Path,
                       help='Arquivo DOCX bilíngue de saída')
    parser.add_argument('--author', '-a', default='Desconhecido',
                       help='Nome do autor')
    parser.add_argument('--keywords', '-k', type=int, default=100,
                       help='Número de palavras-chave (padrão: 100)')
    parser.add_argument('--batch', action='store_true',
                       help='Processa todos em translated/')

    args = parser.parse_args()

    logger.info("="*70)
    logger.info("PIPELINE 2 - PROCESSADOR BILÍNGUE")
    logger.info(f"Palavras-chave: {args.keywords}")
    logger.info("="*70)

    # Modo arquivo único
    if args.input:
        if not args.input.exists():
            logger.error(f"Arquivo não encontrado: {args.input}")
            return 1

        output = args.output or args.input.with_suffix('.docx')
        book_name = args.input.stem.replace('_pt', '')

        success = process_bilingual(
            args.input,
            output,
            args.author,
            book_name,
            args.keywords
        )

        return 0 if success else 1

    # Modo batch
    if args.batch:
        from config.settings import TRANSLATED_DIR

        logger.info(f"Modo batch: {TRANSLATED_DIR}")

        if not TRANSLATED_DIR.exists():
            logger.error(f"Diretório não existe: {TRANSLATED_DIR}")
            return 1

        # Busca arquivos traduzidos
        pt_files = list(TRANSLATED_DIR.rglob("*_pt.txt"))

        if not pt_files:
            logger.warning("Nenhum arquivo traduzido encontrado (*_pt.txt)")
            return 0

        logger.info(f"Encontrados {len(pt_files)} arquivos")

        success_count = 0
        fail_count = 0

        output_base = OUTPUT_DOCX_DIR / "pipeline2"

        for txt_file in sorted(pt_files):
            book_name = txt_file.stem.replace('_pt', '')
            author_name = txt_file.parent.name if txt_file.parent != TRANSLATED_DIR else "Desconhecido"

            output_dir = output_base / author_name
            output_file = output_dir / f"{book_name}_pt_bilingual.docx"

            # Pula se já existe
            if output_file.exists():
                logger.info(f"Já existe: {output_file.name}")
                continue

            logger.info(f"\n--- Processando: {book_name} ({author_name}) ---")

            ok = process_bilingual(
                txt_file,
                output_file,
                author_name,
                book_name,
                args.keywords
            )

            if ok:
                success_count += 1
            else:
                fail_count += 1

        logger.info("\n" + "="*70)
        logger.info(f"RESUMO: Sucesso={success_count}, Falhas={fail_count}")
        logger.info("="*70)

        return 0 if fail_count == 0 else 1

    parser.print_help()
    return 0

if __name__ == "__main__":
    sys.exit(main())
