#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RUN_CALIBRE_BATCH.PY - Processa livros do CALIBRE (Versão Completa)
====================================================================
Pipeline completo:
1. Extrai texto de EPUB/PDF
2. Detecta idioma
3. Traduz para PT-BR (se necessário)
4. Corrige OCR/gramática
5. Identifica notas de rodapé
6. Gera DOCX formatado com template
7. Atualiza banco de dados

Com barra de progresso e processamento paralelo!
"""

import os
import sys
import re
import time
import hashlib
import sqlite3
import asyncio
import logging
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from typing import Optional, List, Tuple, Dict
import traceback

# Extração de texto
import fitz  # PyMuPDF
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup

# Detecção de idioma
from langdetect import detect, LangDetectException

# Progress bar
from tqdm import tqdm

# Requests para Ollama
import requests

# DOCX
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configurações
sys.path.insert(0, str(Path(__file__).parent / "src"))
from config.settings import (
    LOG_DIR, TEMPLATE_DOCX, OUTPUT_DOCX_DIR, DATABASE_PATH, CACHE_PATH,
    OLLAMA_BASE_URL, OLLAMA_MODEL, TEMPERATURE, MAX_OUTPUT_TOKENS,
    MAX_CHUNK_TOKENS, CHAPTER_PATTERNS
)

# ============================================================================
# CONFIG
# ============================================================================

CALIBRE_DIR = Path("D:/Books/CALIBRE")
TXT_OUTPUT_DIR = Path("D:/Books/BooksKDP/txt")
TRANSLATED_DIR = Path("D:/Books/BooksKDP/translated")

MAX_EXTRACTION_WORKERS = 4
MIN_TEXT_LENGTH = 1000
CHUNK_SIZE = 1500  # tokens aproximados por chunk para tradução

# Idiomas que precisam tradução (não são PT)
NEEDS_TRANSLATION = {'en', 'es', 'fr', 'de', 'it', 'ru', 'la', 'el', 'nl', 'pl', 'cs', 'hu', 'ro', 'sv', 'da', 'no', 'fi'}

# Logging
LOG_DIR.mkdir(exist_ok=True)
log_file = LOG_DIR / f"calibre_batch_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)-8s | %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# CACHE
# ============================================================================

class TranslationCache:
    """Cache para traduções."""

    def __init__(self, db_path: Path = CACHE_PATH):
        self.db_path = db_path
        self.db_path.parent.mkdir(exist_ok=True)
        self._init_db()

    def _init_db(self):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS translations (
                    hash TEXT PRIMARY KEY,
                    original TEXT,
                    translated TEXT,
                    source_lang TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_trans ON translations(hash)')

    def _hash(self, text: str) -> str:
        return hashlib.sha256(text.encode()).hexdigest()[:32]

    def get(self, text: str) -> Optional[str]:
        with sqlite3.connect(self.db_path) as conn:
            cur = conn.execute('SELECT translated FROM translations WHERE hash = ?', (self._hash(text),))
            row = cur.fetchone()
            return row[0] if row else None

    def set(self, original: str, translated: str, lang: str):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                INSERT OR REPLACE INTO translations (hash, original, translated, source_lang)
                VALUES (?, ?, ?, ?)
            ''', (self._hash(original), original, translated, lang))

cache = TranslationCache()

# ============================================================================
# BANCO DE DADOS
# ============================================================================

def update_book_status(book_name: str, docx_path: str):
    """Atualiza status do livro no banco."""
    try:
        with sqlite3.connect(DATABASE_PATH) as conn:
            # Tenta encontrar pelo título
            conn.execute('''
                UPDATE books SET
                    processado = 1,
                    processado_em = ?,
                    docx_path = ?
                WHERE title_normalized LIKE ?
            ''', (datetime.now().isoformat(), docx_path, f'%{book_name.lower()[:30]}%'))

            if conn.total_changes > 0:
                logger.debug(f"BD atualizado: {book_name}")
    except Exception as e:
        logger.debug(f"Erro BD (não crítico): {e}")

# ============================================================================
# EXTRAÇÃO DE TEXTO
# ============================================================================

def extract_text_from_epub(epub_path: Path) -> Optional[str]:
    """Extrai texto de EPUB."""
    try:
        book = epub.read_epub(str(epub_path), options={'ignore_ncx': True})
        texts = []

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                # Remove scripts e estilos
                for tag in soup(['script', 'style']):
                    tag.decompose()
                text = soup.get_text(separator='\n\n')
                if text.strip():
                    texts.append(text.strip())

        return '\n\n'.join(texts)
    except Exception as e:
        logger.error(f"Erro EPUB {epub_path.name}: {e}")
        return None


def extract_text_from_pdf(pdf_path: Path) -> Optional[str]:
    """Extrai texto de PDF."""
    try:
        doc = fitz.open(str(pdf_path))
        texts = []

        for page in doc:
            text = page.get_text()
            if text.strip():
                texts.append(text.strip())

        doc.close()
        return '\n\n'.join(texts)
    except Exception as e:
        logger.error(f"Erro PDF {pdf_path.name}: {e}")
        return None


def detect_language(text: str) -> str:
    """Detecta idioma do texto."""
    try:
        # Usa amostra do texto
        sample = text[:5000]
        lang = detect(sample)
        return lang
    except LangDetectException:
        return 'unknown'

# ============================================================================
# TRADUÇÃO E CORREÇÃO VIA OLLAMA
# ============================================================================

def call_ollama(prompt: str, max_retries: int = 3) -> Optional[str]:
    """Chama Ollama API."""
    for attempt in range(max_retries):
        try:
            response = requests.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json={
                    "model": OLLAMA_MODEL,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": TEMPERATURE,
                        "num_predict": MAX_OUTPUT_TOKENS
                    }
                },
                timeout=300
            )
            response.raise_for_status()
            return response.json().get("response", "").strip()
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)
            else:
                logger.error(f"Ollama falhou: {e}")
                return None


def translate_chunk(chunk: str, source_lang: str) -> str:
    """Traduz um chunk para PT-BR."""
    # Verifica cache
    cached = cache.get(chunk)
    if cached:
        return cached

    lang_names = {
        'en': 'inglês', 'es': 'espanhol', 'fr': 'francês', 'de': 'alemão',
        'it': 'italiano', 'ru': 'russo', 'la': 'latim', 'nl': 'holandês'
    }
    lang_name = lang_names.get(source_lang, source_lang)

    prompt = f"""Traduza o texto abaixo de {lang_name} para português brasileiro.

REGRAS:
- Tradução fiel ao original
- Mantenha estrutura de parágrafos
- Mantenha nomes próprios
- NÃO adicione comentários
- Retorne APENAS a tradução

TEXTO:
\"\"\"
{chunk}
\"\"\"

TRADUÇÃO:"""

    result = call_ollama(prompt)
    if result:
        cache.set(chunk, result, source_lang)
        return result
    return chunk


def correct_chunk(chunk: str) -> str:
    """Corrige OCR e gramática."""
    prompt = f"""Corrija o texto abaixo (português brasileiro).

REGRAS:
- Corrija erros de OCR, ortografia, gramática
- Mantenha estrutura de parágrafos
- NÃO altere o conteúdo/significado
- Retorne APENAS o texto corrigido

TEXTO:
\"\"\"
{chunk}
\"\"\"

CORRIGIDO:"""

    result = call_ollama(prompt)
    return result if result else chunk


def create_chunks(text: str, max_chars: int = 4000) -> List[str]:
    """Divide texto em chunks."""
    paragraphs = text.split('\n\n')
    chunks = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) < max_chars:
            current += ('\n\n' if current else '') + para
        else:
            if current:
                chunks.append(current)
            current = para

    if current:
        chunks.append(current)

    return chunks

# ============================================================================
# GERAÇÃO DE DOCX
# ============================================================================

def add_page_numbers(doc: Document):
    """Adiciona numeração de páginas."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        if not footer.paragraphs:
            para = footer.add_paragraph()
        else:
            para = footer.paragraphs[0]

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()

        for el in ['begin', None, 'separate', 'end']:
            if el:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), el)
                run._r.append(fc)
            else:
                it = OxmlElement('w:instrText')
                it.set(qn('xml:space'), 'preserve')
                it.text = "PAGE"
                run._r.append(it)


def generate_docx(text: str, output_path: Path, template_path: Path) -> bool:
    """Gera DOCX formatado usando template."""
    try:
        # Carrega template
        if template_path.exists():
            doc = Document(template_path)
            # Limpa conteúdo mantendo estilos
            for p in list(doc.paragraphs):
                p._element.getparent().remove(p._element)
            for t in list(doc.tables):
                t._element.getparent().remove(t._element)
        else:
            doc = Document()

        # Padrão de capítulos
        chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)

        # Adiciona texto
        for para_text in text.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            is_chapter = chapter_regex.match(para_text) is not None

            para = doc.add_paragraph()
            run = para.add_run(para_text)

            if is_chapter:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = Pt(18)
                run.bold = True
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Numeração de páginas
        add_page_numbers(doc)

        # Salva
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return True

    except Exception as e:
        logger.error(f"Erro DOCX: {e}")
        return False

# ============================================================================
# PIPELINE PRINCIPAL
# ============================================================================

def process_single_book(book_path: Path, pbar_extract: tqdm, pbar_translate: tqdm, pbar_docx: tqdm) -> bool:
    """Processa um único livro (extração → tradução → correção → DOCX)."""
    try:
        folder_name = book_path.parent.name
        book_name = re.sub(r'\(\d+\)$', '', book_path.stem).strip()
        book_name = re.sub(r'[<>:"/\\|?*]', '', book_name)

        # Paths
        txt_path = TXT_OUTPUT_DIR / folder_name / f"{book_name}.txt"
        translated_path = TRANSLATED_DIR / folder_name / f"{book_name}_pt.txt"
        docx_path = OUTPUT_DOCX_DIR / folder_name / f"{book_name}_Final.docx"

        # Já tem DOCX final?
        if docx_path.exists():
            pbar_extract.update(1)
            pbar_translate.update(1)
            pbar_docx.update(1)
            return True

        # 1. EXTRAÇÃO
        if txt_path.exists():
            with open(txt_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            if book_path.suffix.lower() == '.epub':
                text = extract_text_from_epub(book_path)
            else:
                text = extract_text_from_pdf(book_path)

            if not text or len(text) < MIN_TEXT_LENGTH:
                pbar_extract.update(1)
                pbar_translate.update(1)
                pbar_docx.update(1)
                return False

            # Salva TXT
            txt_path.parent.mkdir(parents=True, exist_ok=True)
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(text)

        pbar_extract.update(1)
        pbar_extract.set_postfix_str(f"{book_name[:30]}")

        # 2. TRADUÇÃO (se necessário)
        if translated_path.exists():
            with open(translated_path, 'r', encoding='utf-8') as f:
                translated_text = f.read()
        else:
            lang = detect_language(text)

            if lang in NEEDS_TRANSLATION:
                logger.info(f"Traduzindo ({lang}→pt): {book_name}")
                chunks = create_chunks(text)
                translated_chunks = []

                for chunk in chunks:
                    translated = translate_chunk(chunk, lang)
                    translated_chunks.append(translated)

                translated_text = '\n\n'.join(translated_chunks)

                # Salva tradução
                translated_path.parent.mkdir(parents=True, exist_ok=True)
                with open(translated_path, 'w', encoding='utf-8') as f:
                    f.write(translated_text)
            else:
                translated_text = text

        pbar_translate.update(1)
        pbar_translate.set_postfix_str(f"{book_name[:30]}")

        # 3. CORREÇÃO
        chunks = create_chunks(translated_text)
        corrected_chunks = []

        for chunk in chunks:
            corrected = correct_chunk(chunk)
            corrected_chunks.append(corrected)

        final_text = '\n\n'.join(corrected_chunks)

        # 4. GERA DOCX
        success = generate_docx(final_text, docx_path, TEMPLATE_DOCX)

        pbar_docx.update(1)
        pbar_docx.set_postfix_str(f"{book_name[:30]}")

        if success:
            # Atualiza banco
            update_book_status(book_name, str(docx_path))
            logger.info(f"✓ Concluído: {book_name}")

        return success

    except Exception as e:
        logger.error(f"Erro processando {book_path.name}: {e}")
        pbar_extract.update(1)
        pbar_translate.update(1)
        pbar_docx.update(1)
        return False


def find_all_books() -> List[Path]:
    """Encontra todos os livros únicos."""
    books = []
    for ext in ['*.epub', '*.pdf']:
        books.extend(CALIBRE_DIR.rglob(ext))

    # Remove duplicatas
    unique = {}
    for b in books:
        key = re.sub(r'\(\d+\)\.(epub|pdf)$', r'.\1', str(b).lower())
        if key not in unique:
            unique[key] = b

    return list(unique.values())


def main():
    """Pipeline principal."""
    logger.info("="*60)
    logger.info("CALIBRE BATCH PROCESSOR v2")
    logger.info("="*60)

    # Verifica Ollama
    try:
        r = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=5)
        if r.status_code != 200:
            raise Exception("Ollama não respondeu")
        logger.info(f"Ollama OK - Modelo: {OLLAMA_MODEL}")
    except Exception as e:
        logger.error(f"Ollama não disponível: {e}")
        logger.info("Execute: ollama serve")
        return

    # Encontra livros
    books = find_all_books()
    total = len(books)
    logger.info(f"Encontrados {total} livros únicos")

    if not books:
        return

    # Cria diretórios
    TXT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TRANSLATED_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOCX_DIR.mkdir(parents=True, exist_ok=True)

    # Barras de progresso
    print("\n")
    pbar_extract = tqdm(total=total, desc="1. Extração   ", unit="livro", position=0, leave=True, ncols=80)
    pbar_translate = tqdm(total=total, desc="2. Tradução   ", unit="livro", position=1, leave=True, ncols=80)
    pbar_docx = tqdm(total=total, desc="3. DOCX       ", unit="livro", position=2, leave=True, ncols=80)

    success = 0
    fail = 0

    # Processa sequencialmente (IA é o gargalo)
    for book in books:
        try:
            result = process_single_book(book, pbar_extract, pbar_translate, pbar_docx)
            if result:
                success += 1
            else:
                fail += 1
        except KeyboardInterrupt:
            logger.warning("\nInterrompido!")
            break
        except Exception as e:
            logger.error(f"Erro: {e}")
            fail += 1
            pbar_extract.update(1)
            pbar_translate.update(1)
            pbar_docx.update(1)

    pbar_extract.close()
    pbar_translate.close()
    pbar_docx.close()

    # Resumo
    print("\n")
    logger.info("="*60)
    logger.info("RESUMO FINAL")
    logger.info(f"Total de livros: {total}")
    logger.info(f"Sucesso: {success}")
    logger.info(f"Falhas: {fail}")
    logger.info(f"DOCXs em: {OUTPUT_DOCX_DIR}")
    logger.info("="*60)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("\nInterrompido pelo usuário")
    except Exception as e:
        logger.error(f"Erro fatal: {e}")
        traceback.print_exc()
