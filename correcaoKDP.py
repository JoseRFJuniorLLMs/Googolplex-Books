# -*- coding: utf-8 -*-
# --- Using Google's Gemini API (gemini-1.5-pro) ---

# Standard Python Libraries
import sys # << Importado para sys.executable >>
from dotenv import load_dotenv
import os
import re
import logging
from tqdm import tqdm
import time
import shutil
import traceback # Para log de erros detalhado
import glob # Para encontrar arquivos .txt
import smtplib # For email
import ssl # For email security
from email.message import EmailMessage # For constructing email
import subprocess # << IMPORTADO PARA CHAMAR O SCRIPT TRADUTOR >>

# Third-party Libraries (ensure installed: pip install python-docx google-generativeai python-dotenv tqdm)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.shared import RGBColor
import google.generativeai as genai # Google Generative AI Library

# === SETUP LOGGING ===
log_dir = "logs"
if not os.path.exists(log_dir): os.makedirs(log_dir)
log_filepath = os.path.join(log_dir, "book_processor_multi_author_mem.log")
# Log para arquivos com CORREÇÃO concluída
PROCESSED_LOG_FILE = os.path.join(log_dir, "processed_books.log")
# Log para arquivos com TRADUÇÃO concluída (NOVO)
TRANSLATED_LOG_FILE = os.path.join(log_dir, "translated_books.log")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(module)s:%(lineno)d - %(funcName)s - %(message)s',
    handlers=[ logging.FileHandler(log_filepath, encoding='utf-8'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
# Email configuration from .env file (CARREGADO AQUI)
EMAIL_SENDER_ADDRESS = os.getenv("EMAIL_SENDER_ADDRESS")
EMAIL_SENDER_APP_PASSWORD = os.getenv("EMAIL_SENDER_APP_PASSWORD") # Use App Password for Gmail
EMAIL_RECIPIENT_ADDRESS = os.getenv("EMAIL_RECIPIENT_ADDRESS", "web2ajax@gmail.com") # Destinatário padrão
EMAIL_SMTP_SERVER = os.getenv("EMAIL_SMTP_SERVER", "smtp.gmail.com") # Default Gmail
EMAIL_SMTP_PORT = int(os.getenv("EMAIL_SMTP_PORT", 587)) # Default Gmail Port (TLS)


# === CONFIGURAÇÕES ===

# -- Diretórios Base --
BASE_INPUT_TXT_DIR = "txt"
BASE_OUTPUT_DOCX_DIR = "docx"
BASE_OUTPUT_TXT_DIR = "txt"

# -- Nomes de Arquivos Base --
# !! AJUSTE SE NECESSÁRIO !!
TEMPLATE_DOCX = "Estrutura.docx" # Template OBRIGATÓRIO

# -- Nomes Base dos Arquivos de Saída (Correção) --
FINAL_DOCX_BASENAME = "Livro_Final_Formatado_Sem_Notas.docx"
FINAL_NUMBERED_TXT_BASENAME = "Livro_Final_Com_Notas_Numeros.txt"
NOTES_TXT_FILE_BASENAME = "notas_rodape.txt"

# -- Nomes Base dos Arquivos de Saída (Tradução - NOVO) --
# O nome exato será gerado dentro do loop principal
TRANSLATED_DOCX_SUFFIX = "-A0.docx"

# -- Configurações da API e Processamento (Correção) --
MODEL_NAME = "gemini-1.5-pro" # Modelo para correção e notas
MAX_CHUNK_TOKENS = 1500 # Aprox. limite de tokens por chunk para API
MAX_OUTPUT_TOKENS = 8192 # Limite de saída do Gemini
TEMPERATURE = 0.5 # Temperatura para a correção

# -- Configurações para o Script Tradutor (NOVO) --
# !! AJUSTE ESTE CAMINHO para onde você salvou o script_tradutor_hibrido.py !!
PATH_TO_TRANSLATOR_SCRIPT = "script_tradutor_hibrido.py"
NUM_WORDS_TO_TRANSLATE = 100 # Quantidade de palavras para o script tradutor processar

# -- Estilos e Padrões (Correção) --
NORMAL_STYLE_NAME = "Normal" # Estilo esperado no template
CHAPTER_PATTERNS = [
    r'^\s*Capítulo \w+', r'^\s*CAPÍTULO \w+', r'^\s*Capítulo \d+',
    r'^\s*CHAPTER \w+', r'^\s*Chapter \d+', r'^\s*LIVRO \w+', r'^\s*PARTE \w+',
    # Adicionar outros padrões se necessário
]
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA===" # Marcador para quebra de página manual no texto
AI_FAILURE_MARKER = "*** FALHA NA IA - TEXTO ORIGINAL ABAIXO ***" # Marcador de falha da API
FORMATTING_ERROR_MARKER = "*** ERRO DE FORMATAÇÃO - TEXTO ORIGINAL ABAIXO ***" # Marcador de erro de formatação

# --- Validação API Key ---
if not GOOGLE_API_KEY:
    logger.error("FATAL: GOOGLE_API_KEY não encontrada nas variáveis de ambiente (.env).")
    exit(1)

# --- Validação Config Email (Informativo) ---
if not EMAIL_SENDER_ADDRESS or not EMAIL_SENDER_APP_PASSWORD or not EMAIL_RECIPIENT_ADDRESS:
    logger.warning("AVISO: Variáveis de ambiente para envio de e-mail (EMAIL_SENDER_ADDRESS, EMAIL_SENDER_APP_PASSWORD, EMAIL_RECIPIENT_ADDRESS) não configuradas ou incompletas no .env. Notificação por e-mail será desativada.")
else:
    logger.info(f"Configuração de e-mail carregada. Notificações serão enviadas de '{EMAIL_SENDER_ADDRESS}' para '{EMAIL_RECIPIENT_ADDRESS}'.")

# --- Setup Gemini Client ---
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    # Configurações de segurança mais permissivas (use com cautela)
    safety_settings_lenient = {
        'HATE': 'BLOCK_NONE', 'HARASSMENT': 'BLOCK_NONE',
        'SEXUAL' : 'BLOCK_NONE', 'DANGEROUS' : 'BLOCK_NONE'
    }
    generation_config = genai.GenerationConfig(
                    temperature=TEMPERATURE,
                    max_output_tokens=MAX_OUTPUT_TOKENS
                )
    gemini_model = genai.GenerativeModel(
        MODEL_NAME,
        safety_settings=safety_settings_lenient,
        generation_config=generation_config
    )
    logger.info(f"Modelo Gemini '{MODEL_NAME}' inicializado com sucesso.")
except Exception as e:
    logger.error(f"FATAL: Falha ao inicializar modelo Gemini ({MODEL_NAME}): {e}")
    logger.error(traceback.format_exc())
    exit(1)

# --- Funções Auxiliares ---

def count_tokens_approx(text):
    """Estima a contagem de tokens (aproximadamente 3 chars/token)."""
    if not text: return 0
    return len(text) // 3

def create_chunks(text, max_tokens, author_name="N/A", book_name="N/A"):
    """Divide o texto em chunks, subdividindo parágrafos grandes."""
    log_prefix = f"[{author_name}/{book_name}]"
    # logger.info(f"{log_prefix} Iniciando criação de chunks. Máx tokens (aprox): {max_tokens}") # Verbose
    chunks = []
    current_chunk = ""
    current_chunk_tokens = 0
    # Divide por parágrafos (duas quebras de linha)
    paragraphs = text.split("\n\n")
    # paragraphs_stripped = [p.strip() for p in paragraphs if p.strip()] # Não usado?

    # logger.info(f"{log_prefix} Texto dividido inicialmente em {len(paragraphs_stripped)} blocos não vazios ('\\n\\n').") # Verbose

    for i, paragraph_text in enumerate(paragraphs):
        # Ignora blocos vazios mas tenta manter o espaçamento
        if not paragraph_text.strip():
            if chunks and chunks[-1].strip():
                 if not chunks[-1].endswith("\n\n"):
                      chunks[-1] += "\n\n"
            continue

        paragraph_tokens = count_tokens_approx(paragraph_text)
        tokens_with_separator = paragraph_tokens + (count_tokens_approx("\n\n") if current_chunk else 0)

        # Se o parágrafo cabe no chunk atual, adiciona
        if current_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + paragraph_text
            current_chunk_tokens = count_tokens_approx(current_chunk)
        # Se não cabe, salva o chunk atual e inicia um novo com o parágrafo
        else:
            if current_chunk: # Salva o chunk anterior se ele não estiver vazio
                chunks.append(current_chunk)
                # logger.debug(f"{log_prefix} Chunk {len(chunks)} salvo (limite atingido). Tokens: {current_chunk_tokens}.") # Verbose
            # O parágrafo atual começa um novo chunk
            current_chunk = paragraph_text
            current_chunk_tokens = paragraph_tokens

            # --- Subdivisão se o *próprio* parágrafo atual for muito grande ---
            if paragraph_tokens > max_tokens:
                logger.warning(f"{log_prefix} Parágrafo {i+1} ({paragraph_tokens} tk) excede limite {max_tokens}. Iniciando SUBDIVISÃO.")
                # Remove o parágrafo grande do current_chunk (que acabamos de adicionar)
                current_chunk = ""
                current_chunk_tokens = 0

                sub_chunks_added_count = 0
                # Tenta dividir por frases, senão por linhas
                sentences = re.split(r'(?<=[.!?])\s+', paragraph_text)
                if len(sentences) <= 1 :
                    sentences = paragraph_text.split('\n') # Fallback para linhas

                current_sub_chunk = ""
                current_sub_chunk_tokens = 0
                for sentence_num, sentence in enumerate(sentences):
                    sentence_clean = sentence.strip()
                    if not sentence_clean: continue

                    sentence_tokens = count_tokens_approx(sentence)
                    tokens_with_sub_separator = sentence_tokens + (count_tokens_approx("\n") if current_sub_chunk else 0)

                    # Se a sentença cabe no sub-chunk atual
                    if current_sub_chunk_tokens + tokens_with_sub_separator <= max_tokens:
                        sub_separator = "\n" if current_sub_chunk else "" # Usa \n dentro do parágrafo subdividido
                        current_sub_chunk += sub_separator + sentence
                        current_sub_chunk_tokens = count_tokens_approx(current_sub_chunk)
                    # Se não cabe, salva o sub-chunk e inicia um novo
                    else:
                        if current_sub_chunk: # Salva o sub-chunk anterior
                            chunks.append(current_sub_chunk)
                            sub_chunks_added_count += 1
                            # logger.debug(f"{log_prefix} Sub-chunk {len(chunks)} salvo (limite sub). Tokens: {current_sub_chunk_tokens}.") # Verbose

                        # Verifica se a *própria* sentença é grande demais
                        if sentence_tokens > max_tokens:
                            chunks.append(sentence) # Adiciona a sentença longa como um chunk próprio
                            sub_chunks_added_count += 1
                            logger.warning(f"{log_prefix}  -> Sentença/Linha {sentence_num+1} ({sentence_tokens} tk) excede limite. Adicionada como chunk individual (PODE FALHAR NA API).")
                            current_sub_chunk = "" # Reseta, pois ela foi adicionada separadamente
                            current_sub_chunk_tokens = 0
                        else:
                            # A sentença não é grande demais, ela inicia o novo sub-chunk
                            current_sub_chunk = sentence
                            current_sub_chunk_tokens = sentence_tokens

                # Salva o último sub-chunk restante
                if current_sub_chunk:
                    chunks.append(current_sub_chunk)
                    sub_chunks_added_count += 1
                    # logger.debug(f"{log_prefix} Último sub-chunk {len(chunks)} salvo (Parág. {i+1}). Tokens: {current_sub_chunk_tokens}.") # Verbose

                # Se não conseguiu subdividir (caso raro), adiciona o parágrafo original
                if sub_chunks_added_count == 0:
                     logger.warning(f"{log_prefix} Parágrafo {i+1} excedeu limite, mas não foi subdividido. Adicionando original como chunk (PODE FALHAR NA API).")
                     chunks.append(paragraph_text)

                # Reseta o chunk principal após lidar com o parágrafo grande
                current_chunk = ""
                current_chunk_tokens = 0
            # --- Fim da Subdivisão ---

    # Adiciona o último chunk que sobrou
    if current_chunk:
        chunks.append(current_chunk)
        # logger.debug(f"{log_prefix} Chunk final {len(chunks)} salvo. Tokens: {current_chunk_tokens}.") # Verbose

    # --- Pós-processamento: Junta chunks pequenos consecutivos ---
    # logger.debug(f"{log_prefix} Iniciando merge de chunks pequenos...") # Verbose
    merged_chunks = []
    temp_chunk = ""
    temp_chunk_tokens = 0
    for i, chunk in enumerate(chunks):
        chunk_tokens = count_tokens_approx(chunk)
        tokens_with_separator = chunk_tokens + (count_tokens_approx("\n\n") if temp_chunk else 0)

        # Se o chunk atual cabe junto com o temporário
        if temp_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if temp_chunk else ""
            temp_chunk += separator + chunk
            temp_chunk_tokens = count_tokens_approx(temp_chunk)
        # Se não cabe, salva o temporário e inicia um novo
        else:
            if temp_chunk: merged_chunks.append(temp_chunk)
            temp_chunk = chunk
            temp_chunk_tokens = chunk_tokens

    # Salva o último chunk temporário
    if temp_chunk:
        merged_chunks.append(temp_chunk)

    final_chunk_count = len(merged_chunks)
    if final_chunk_count < len(chunks):
         logger.info(f"{log_prefix} Merge concluído. De {len(chunks)} para {final_chunk_count} chunks.")
    # else: logger.debug(f"{log_prefix} Merge não alterou número de chunks ({final_chunk_count}).") # Verbose

    logger.info(f"{log_prefix} ✅ Chunking concluído. {final_chunk_count} chunks finais.")
    return merged_chunks

def _call_gemini_api(model, prompt_text, chunk_for_log, author_name="N/A", book_name="N/A"):
    """Função interna para chamar a API Gemini com retries e tratamento de erro."""
    log_prefix = f"[{author_name}/{book_name}]"
    max_retries = 5
    base_wait_time = 5 # Segundos iniciais de espera
    log_chunk_preview = chunk_for_log[:150].replace('\n', '\\n') + ('...' if len(chunk_for_log) > 150 else '')

    for attempt in range(max_retries):
        # logger.info(f"{log_prefix} Chamando API (Tentativa {attempt + 1}/{max_retries}). Chunk (início): '{log_chunk_preview}'") # Verbose
        try:
            # Usa a configuração de geração definida ao criar o modelo
            response = model.generate_content(prompt_text)

            # Verifica bloqueio de prompt primeiro
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback and \
               hasattr(response.prompt_feedback, 'block_reason') and response.prompt_feedback.block_reason:
                block_reason = response.prompt_feedback.block_reason.name
                logger.error(f"{log_prefix} API BLOQUEOU O PROMPT (Tentativa {attempt + 1}). Razão: {block_reason}. Chunk: '{log_chunk_preview}'")
                return None # Falha imediata se o prompt for bloqueado

            # Verifica se há candidatos na resposta
            if not response.candidates:
                 logger.error(f"{log_prefix} API retornou SEM CANDIDATOS (Tentativa {attempt + 1}). Resposta: {response}. Chunk: '{log_chunk_preview}'")
                 # Continua tentando, pode ser erro temporário
            else:
                 # Processa o primeiro candidato (geralmente o único)
                 try:
                    candidate = response.candidates[0]
                    finish_reason = candidate.finish_reason.name if hasattr(candidate, 'finish_reason') and candidate.finish_reason else "FINISH_REASON_UNKNOWN"
                    safety_ratings = [(r.category.name, r.probability.name) for r in candidate.safety_ratings] if candidate.safety_ratings else "N/A"
                    # logger.debug(f"{log_prefix} API Call OK (Tentativa {attempt + 1}). Finish: {finish_reason}. Safety: {safety_ratings}") # Verbose

                    # Avisa sobre motivos de término não ideais
                    if finish_reason == "STOP": pass # OK
                    elif finish_reason == "MAX_TOKENS": logger.warning(f"{log_prefix} API TRUNCOU resposta devido a MAX_OUTPUT_TOKENS.")
                    elif finish_reason == "SAFETY": logger.warning(f"{log_prefix} API interrompeu resposta devido a SAFETY. Ratings: {safety_ratings}.")
                    elif finish_reason == "RECITATION": logger.warning(f"{log_prefix} API interrompeu resposta devido a RECITATION.")
                    elif finish_reason == "OTHER": logger.warning(f"{log_prefix} API interrompeu resposta por OUTRA RAZÃO.")
                    else: logger.warning(f"{log_prefix} API retornou com finish_reason inesperado: {finish_reason}.")

                    # Tenta extrair o texto da resposta
                    result_text = ""
                    if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                        text_parts = [part.text for part in candidate.content.parts if hasattr(part, 'text')]
                        if text_parts:
                            result_text = "".join(text_parts).strip()
                        # else: logger.warning(f"{log_prefix} Resposta API tem 'parts' mas não foi possível extrair texto (Tentativa {attempt+1}). Parts: {candidate.content.parts}") # Verbose
                    elif hasattr(response, 'text') and response.text: # Fallback para modelo mais antigo?
                        result_text = response.text.strip()
                        # logger.debug(f"{log_prefix} Texto extraído via response.text (fallback).") # Verbose

                    # Se conseguiu extrair texto, retorna
                    if result_text:
                        # logger.debug(f"{log_prefix} Texto API recebido (100 chars): '{result_text[:100].replace('\n', '\\n')}...'") # Verbose
                        # Checagem simples: resposta muito curta comparada à entrada?
                        if len(result_text) < len(chunk_for_log) * 0.1 and len(chunk_for_log) > 100:
                            logger.warning(f"{log_prefix} Resposta da API parece muito curta. Input len: {len(chunk_for_log)}, Output len: {len(result_text)}.")
                        return result_text # SUCESSO!
                    else:
                         # Se não conseguiu extrair texto, mesmo com candidato, loga e tenta de novo
                         logger.warning(f"{log_prefix} Resposta API não continha texto utilizável (Tentativa {attempt+1}), mesmo com candidato. Finish Reason: {finish_reason}.")

                 except Exception as e_details:
                    # Erro ao processar a resposta bem-sucedida
                    logger.error(f"{log_prefix} Erro ao extrair detalhes/texto da resposta API (Tentativa {attempt+1}): {e_details} - Resposta Crua: {response}")
                    logger.error(traceback.format_exc())
                    # Continua tentando

            # --- Espera Exponencial com Jitter antes de tentar de novo ---
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time) # Jitter
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                 # Se chegou aqui, esgotou as tentativas sem sucesso
                 logger.error(f"{log_prefix} Falha final na API após {max_retries} tentativas para o chunk: '{log_chunk_preview}'")
                 return None # Falha após todas as tentativas

        except Exception as e: # Captura erros durante a chamada da API em si (conexão, etc.)
            logger.warning(f"{log_prefix} Erro durante a chamada da API ({model.model_name}) (Tentativa {attempt + 1}/{max_retries}): {e}")
            logger.error(traceback.format_exc())
            # Tratamento específico para erros comuns
            if "RESOURCE_EXHAUSTED" in str(e) or "429" in str(e):
                 logger.warning(f"{log_prefix} Erro de cota (RESOURCE_EXHAUSTED / 429). Aumentando espera.")
                 base_wait_time = max(15, base_wait_time) # Aumenta espera base para erros de cota
            elif "Internal error encountered." in str(e) or "500" in str(e):
                 logger.warning(f"{log_prefix} Erro interno do servidor (500). Tentando novamente.")
                 # Mantém espera normal
            # Aplica espera exponencial para outros erros também
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                # Esgotou tentativas após erro na chamada
                logger.error(f"{log_prefix} Falha final na API após {max_retries} tentativas (erro na chamada) para o chunk: '{log_chunk_preview}'")
                return None # Falha após todas as tentativas

    # Se o loop terminar sem retornar sucesso ou None explicitamente (improvável)
    logger.error(f"{log_prefix} Loop de tentativas da API concluído sem sucesso explícito para o chunk: '{log_chunk_preview}'")
    return None

def format_with_ai_correction_only(model, chunk, author_name, book_name, is_first_chunk=False):
    """Chama a API Gemini focando APENAS na correção de OCR/gramática."""
    log_prefix = f"[{author_name}/{book_name}]"
    context_start = "Você está formatando o início de um livro." if is_first_chunk else "Você está continuando a formatação de um texto de livro existente."
    # Exemplo de erros de OCR (pode ser expandido)
    ocr_errors_examples = """
        * **Troca de letras similares:** 'rn' vs 'm', 'c' vs 'e', 't' vs 'f', 'l' vs 'i', 'I' vs 'l', 'O' vs '0', 'S' vs '5', 'B' vs '8'.
        * **Hífens indevidos:** Palavras quebradas incorretamente no meio ou hífens extras.
        * **Hífens ausentes:** Palavras que deveriam ser hifenizadas (ex: "guarda-chuva") aparecem juntas ou separadas.
        * **Espaços ausentes ou extras:** Palavras coladas ("onomundo") ou espaços excessivos.
        * **Pontuação incorreta:** Pontos finais trocados por vírgulas, pontos de interrogação/exclamação mal interpretados.
        * **Acentuação:** Falta de acentos (ex: 'e' vs 'é', 'a' vs 'à'), acentos incorretos (crase onde não deve) ou caracteres estranhos no lugar de acentos.
        * **Letras duplicadas ou ausentes:** "caaasa" ou "casaa" em vez de "casa".
        * **Confusão maiúsculas/minúsculas:** Nomes próprios em minúsculas, inícios de frase em minúsculas.
        * **Caracteres especiais/ruído:** Símbolos aleatórios '%', '#', '@' inseridos no texto.
        * **Quebras de linha estranhas:** Parágrafos divididos no meio sem motivo aparente. Preserve as quebras de parágrafo intencionais (duas quebras de linha).
    """
    chunk_prompt = f"""
{context_start} Você é um editor literário proficiente em português do Brasil. Sua tarefa é CORRIGIR e FORMATAR o fragmento de texto a seguir, que pertence a um livro do autor {author_name}.

**CONTEXTO IMPORTANTE:** Este texto provavelmente foi extraído via OCR de um PDF e pode conter erros de reconhecimento, digitação e gramática. O objetivo principal é obter um texto LIMPO e CORRETO em português do Brasil padrão, mantendo a estrutura e o significado originais.

**SIGA RIGOROSAMENTE ESTAS REGRAS:**

1.  **Correção Profunda:** Corrija TODOS os erros gramaticais, ortográficos, de pontuação, acentuação e concordância verbal/nominal. Use o português do Brasil como referência. FOQUE em erros comuns de OCR como os listados abaixo.
2.  **Estilo e Tom:** Mantenha o estilo literário e o tom do texto original do autor {author_name}. Seja claro, fluido e envolvente. NÃO altere o significado, a voz ou a intenção do autor.
3.  **Fidelidade Estrutural:** MANTENHA a estrutura de parágrafos original. Parágrafos são geralmente separados por UMA linha em branco (duas quebras de linha `\\n\\n`). NÃO junte parágrafos que estavam separados. NÃO divida parágrafos desnecessariamente.
4.  **Sem Adições/Remoções:** NÃO omita frases ou informações. NÃO adicione conteúdo, introduções, resumos, conclusões ou opiniões que não estavam no fragmento original. SEJA ESTRITAMENTE FIEL AO CONTEÚDO.
5.  **Marcadores de Capítulo/Quebra:** Se encontrar marcadores como 'Capítulo X', '***', '---', etc., no início de um parágrafo, MANTENHA-OS EXATAMENTE como estão, naquele parágrafo específico. NÃO adicione ou remova esses marcadores.
6.  **Quebra de Página:** Se o marcador '{PAGE_BREAK_MARKER}' aparecer, MANTENHA-O EXATAMENTE onde está, em sua própria linha, sem texto antes ou depois na mesma linha.
7.  **Erros Comuns de OCR (FOCO ESPECIAL):** Preste atenção e corrija diligentemente:
    {ocr_errors_examples}
8.  **Formato de Saída:** Retorne APENAS o texto corrigido e formatado. Use parágrafos separados por duas quebras de linha (`\\n\\n`). NÃO use NENHUMA formatação especial como Markdown (`*`, `#`, `_`), HTML, etc. Retorne TEXTO PURO. Não inclua comentários sobre o que você fez, apenas o texto resultante.

**Texto do fragmento para processar (pode conter erros):**
\"\"\"
{chunk}
\"\"\"

**Lembre-se: O resultado deve ser APENAS o texto corrigido.**
"""
    # logger.debug(f"{log_prefix} Enviando chunk para CORREÇÃO (API: {model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk") # Verbose
    return _call_gemini_api(model, chunk_prompt, chunk, author_name, book_name)


def format_with_ai_footnote_only(model, chunk, author_name, book_name):
    """Chama a API Gemini focando APENAS na identificação de notas."""
    log_prefix = f"[{author_name}/{book_name}]"
    chunk_prompt = f"""
Você é um assistente de edição trabalhando no texto do autor {author_name}. Sua tarefa é analisar o fragmento de texto A SEGUIR, que JÁ FOI CORRIGIDO no passo anterior, e APENAS inserir marcadores para potenciais notas de rodapé onde estritamente necessário.

**REGRAS IMPORTANTES:**

1.  **NÃO ALTERE O TEXTO CORRIGIDO:** Não faça correções, não mude palavras, não reestruture frases. Apenas insira os marcadores.
2.  **MARCADORES DE NOTA:** Insira marcadores APENAS nos seguintes casos:
    * **Termos em Idioma Estrangeiro (não comuns):** Imediatamente APÓS uma palavra ou frase curta em latim, francês, inglês, etc., que não seja de uso corrente em português, insira: `[NOTA_IDIOMA:palavra_original][CONTEUDO_NOTA:Tradução ou breve explicação]`. Exemplo: "...uma certa *joie de vivre*[NOTA_IDIOMA:joie de vivre][CONTEUDO_NOTA:Alegria de viver (francês)]..."
    * **Citações/Referências:** APÓS uma citação direta curta ou uma referência bibliográfica no texto (ex: (Autor, Ano)), insira: `[NOTA_CITACAO:Texto citado ou referência][CONTEUDO_NOTA:Referência bibliográfica completa ou fonte, se conhecida ou inferível]`. Exemplo: "...como disse Foucault (1975)[NOTA_CITACAO:Foucault (1975)][CONTEUDO_NOTA:FOUCAULT, Michel. Vigiar e Punir. 1975.], a disciplina..."
    * **Nomes Próprios (contexto essencial):** APÓS um nome de pessoa, local ou evento histórico POUCO CONHECIDO que SEJA ESSENCIAL contextualizar brevemente para a compreensão do trecho, insira: `[NOTA_NOME:Nome Mencionado][CONTEUDO_NOTA:Breve identificação (datas, relevância)]`. Use com MODERAÇÃO. Exemplo: "...influenciado por Kropotkin[NOTA_NOME:Kropotkin][CONTEUDO_NOTA:Piotr Kropotkin (1842-1921), anarquista russo.]..."
    * **Termos Técnicos/Jargão (essencial):** APÓS um termo técnico MUITO específico de uma área, cuja definição SEJA INDISPENSÁVEL para o leitor geral entender o argumento naquele ponto, insira: `[NOTA_TERMO:Termo Técnico][CONTEUDO_NOTA:Definição concisa]`. Use com MUITA MODERAÇÃO. Exemplo: "...aplicando a análise de isotopias[NOTA_TERMO:Isotopias][CONTEUDO_NOTA:Na semiótica greimasiana, recorrência de categorias sêmicas que garante a homogeneidade de um discurso.]..."
3.  **FORMATO DOS MARCADORES:** Use EXATAMENTE `[NOTA_TIPO:Referência]` seguido IMEDIATAMENTE por `[CONTEUDO_NOTA:Explicação]`. Não adicione espaços entre eles. Não use outros formatos.
4.  **CRITÉRIO:** Seja conservador. Adicione notas apenas se a informação for realmente útil e provavelmente desconhecida para um leitor culto médio. É MELHOR ERRAR POR NÃO ADICIONAR do que por adicionar excessivamente. NÃO adicione notas para termos comuns, nomes famosos ou citações óbvias.
5.  **NÃO INVENTE CONTEÚDO:** O `[CONTEUDO_NOTA:...]` deve ser uma tradução direta, uma referência óbvia, ou uma contextualização muito breve e factual, se possível inferida do próprio texto ou conhecimento geral básico. NÃO pesquise externamente para criar notas complexas. Se não souber o conteúdo, NÃO insira a nota.
6.  **SAÍDA:** Retorne APENAS o texto original (do input) com os marcadores inseridos nos locais exatos. Mantenha a estrutura de parágrafos (`\\n\\n`). Não adicione NENHUM outro texto, comentário ou explicação.

**Texto JÁ CORRIGIDO para analisar e inserir marcadores de nota:**
\"\"\"
{chunk}
\"\"\"

**Lembre-se: NÃO altere o texto, apenas insira os marcadores `[NOTA_...][CONTEUDO_NOTA:...]` quando apropriado.**
"""
    # logger.debug(f"{log_prefix} Enviando chunk para IDENTIFICAÇÃO DE NOTAS (API: {model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk") # Verbose
    return _call_gemini_api(model, chunk_prompt, chunk, author_name, book_name)


# --- FUNÇÕES DE PROCESSAMENTO DOS PASSOS ---

def apply_formatting_pass1(doc, formatted_chunk_text, normal_style_name, chapter_patterns, corrected_text_list, author_name, book_name):
    """
    Aplica formatação ao DOCX (Passo 1 - sem notas) e coleta texto para retornar.
    """
    log_prefix = f"[{author_name}/{book_name}]"
    if not formatted_chunk_text or not formatted_chunk_text.strip():
        # logger.warning(f"{log_prefix} Chunk formatado vazio ou apenas espaços recebido (Passo 1). Pulando.") # Verbose
        return

    # Coleta texto para o próximo passo (identificação de notas)
    plain_text_for_list = formatted_chunk_text.replace(PAGE_BREAK_MARKER, "\n\n").strip()
    if plain_text_for_list:
        corrected_text_list.append(plain_text_for_list)
    # else: # Não adiciona chunks que só tinham quebra de página ou eram vazios
        # if formatted_chunk_text.strip() == PAGE_BREAK_MARKER:
             # logger.debug(f"{log_prefix} Chunk continha apenas marcador de página, não adicionado à lista de texto.") # Verbose
        # else:
             # logger.warning(f"{log_prefix} Texto formatado resultou em vazio. Original: '{formatted_chunk_text[:50]}...'") # Verbose

    # --- Aplica formatação ao DOCX ---
    normal_style = None
    try: # Busca o estilo 'Normal' (a existência já foi checada uma vez)
        if normal_style_name in doc.styles:
             normal_style = doc.styles[normal_style_name]
    except Exception as e_style:
         logger.error(f"{log_prefix} Erro ao acessar estilo '{normal_style_name}': {e_style}.")

    chapter_regex = re.compile('|'.join(chapter_patterns), re.IGNORECASE)
    # Processa partes separadas por quebra de página
    parts = formatted_chunk_text.split(PAGE_BREAK_MARKER)
    # Verifica se já há conteúdo antes de adicionar quebra (evita quebra no início)
    content_present_before = any(p.text.strip() for p in doc.paragraphs)

    for part_index, part in enumerate(parts):
        part_clean = part.strip()

        # Adiciona quebra de página ANTES da nova parte (exceto a primeira)
        if part_index > 0:
             # Evita adicionar quebras de página duplicadas
             last_para_is_page_break = False
             if doc.paragraphs:
                 last_p = doc.paragraphs[-1]
                 # Verifica se o último parágrafo é vazio e contém uma quebra de página
                 if not last_p.text.strip() and any(run.text and '\f' in run.text for run in last_p.runs):
                     last_para_is_page_break = True
             if not last_para_is_page_break:
                 # logger.debug(f"{log_prefix} Adicionando quebra de página ao DOCX (antes da parte {part_index + 1}).") # Verbose
                 doc.add_page_break()
             # else: logger.debug(f"{log_prefix} Quebra de página omitida (duplicada).") # Verbose

        # Pula partes vazias (geralmente após uma quebra de página)
        if not part_clean:
            if part_index > 0 : content_present_before = True # Marca que houve conteúdo antes
            continue

        # Processa parágrafos dentro da parte
        paragraphs_in_part = part_clean.split("\n\n")
        for paragraph_text in paragraphs_in_part:
            paragraph_text_clean = paragraph_text.strip()
            # Lida com parágrafos intencionalmente vazios (para espaçamento)
            if not paragraph_text_clean:
                # Só adiciona se o parágrafo anterior não era vazio
                if doc.paragraphs and doc.paragraphs[-1].text.strip():
                     p = doc.add_paragraph()
                     if normal_style: p.style = normal_style # Aplica estilo se disponível
                continue

            # Detecta marcadores especiais / capítulos
            is_ai_failure_marker = paragraph_text_clean.startswith(AI_FAILURE_MARKER)
            is_formatting_error_marker = paragraph_text_clean.startswith(FORMATTING_ERROR_MARKER)
            is_chapter = not is_ai_failure_marker and not is_formatting_error_marker and chapter_regex.match(paragraph_text_clean) is not None

            # Adiciona parágrafo e texto
            p = doc.add_paragraph()
            run = p.add_run(paragraph_text_clean)
            content_present_before = True # Marca que agora há conteúdo

            # Aplica formatação/estilos específicos
            try:
                if is_chapter:
                    # Formatação específica para capítulos (sobrescreve template)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.name = 'French Script MT' # Exemplo
                    run.font.size = Pt(48)             # Exemplo
                    run.bold = False
                    # logger.debug(f"{log_prefix} Aplicada formatação específica de capítulo.") # Verbose
                elif is_ai_failure_marker or is_formatting_error_marker:
                    # Formatação para marcadores de erro
                    if normal_style: p.style = normal_style # Usa Normal como base
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Vermelho
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # logger.debug(f"{log_prefix} Aplicada formatação de marcador de erro.") # Verbose
                else: # Parágrafo normal
                    # Aplica estilo 'Normal' SOMENTE se encontrado no template
                    if normal_style:
                        p.style = normal_style
                    # Se não, depende do padrão do documento (definido no template)
            except Exception as e_apply_style:
                 logger.error(f"{log_prefix} Erro ao aplicar estilo/formatação: {e_apply_style}. Texto: '{paragraph_text_clean[:50]}...'")
                 # Continua sem a formatação específica em caso de erro


def run_correction_pass(model, input_txt_path, template_docx_path, output_docx_path, author_name, book_name):
    """
    Executa o Passo 1: Corrige texto e gera DOCX base usando template.
    Retorna: (bool, str | None) -> (success_status, corrected_text_content)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 1: Correção e Geração DOCX ---")
    logger.info(f"{log_prefix} Lendo texto original de: {input_txt_path}")
    try:
        with open(input_txt_path, "r", encoding="utf-8") as f: texto_original = f.read()
        logger.info(f"{log_prefix} Entrada '{os.path.basename(input_txt_path)}' carregada ({len(texto_original)} chars).")
    except FileNotFoundError:
        logger.error(f"{log_prefix} FATAL: Arquivo de entrada '{input_txt_path}' não encontrado."); return (False, None)
    except Exception as e:
        logger.error(f"{log_prefix} FATAL ao ler entrada '{input_txt_path}': {e}")
        logger.error(traceback.format_exc()); return (False, None)

    # Garante que diretório de saída existe
    output_docx_dir = os.path.dirname(output_docx_path)
    os.makedirs(output_docx_dir, exist_ok=True) # Cria se não existir

    # Divide o texto original em chunks
    logger.info(f"{log_prefix} Dividindo texto original em chunks...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS, author_name, book_name)
    if not text_chunks:
        logger.error(f"{log_prefix} Nenhum chunk gerado do texto original. Abortando Passo 1."); return (False, None)
    # Removido log incorreto que dizia 1 chunk
    # logger.info(f"{log_prefix} Texto original dividido em {len(text_chunks)} chunks.") # Correto

    doc = None
    normal_style_exists = False # Flag para checar estilo uma vez
    logger.info(f"{log_prefix} Preparando documento DOCX para: {os.path.basename(output_docx_path)} usando template OBRIGATÓRIO.")
    try:
        # --- Backup do DOCX existente (se houver) ---
        if os.path.exists(output_docx_path):
             backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
             backup_docx_path = os.path.join(output_docx_dir, f"backup_{os.path.splitext(os.path.basename(output_docx_path))[0]}_{backup_timestamp}.docx")
             try:
                 shutil.copy2(output_docx_path, backup_docx_path)
                 logger.info(f"{log_prefix} Backup do DOCX anterior criado: {os.path.basename(backup_docx_path)}")
             except Exception as e_bkp:
                 logger.warning(f"{log_prefix} Falha ao criar backup de '{os.path.basename(output_docx_path)}': {e_bkp}")

        # --- Carrega Template OBRIGATÓRIO ---
        if not os.path.exists(template_docx_path):
            logger.error(f"{log_prefix} FATAL: Template OBRIGATÓRIO '{template_docx_path}' não encontrado. Abortando.")
            return (False, None)

        try:
            doc = Document(template_docx_path)
            logger.info(f"{log_prefix} Template '{os.path.basename(template_docx_path)}' carregado. Limpando conteúdo existente...")
            # Limpa conteúdo (parágrafos, tabelas) do corpo do template
            # Preserva estilos, cabeçalhos, rodapés, setup de página, etc.
            for para in doc.paragraphs: p_element = para._element; p_element.getparent().remove(p_element)
            for table in doc.tables: tbl_element = table._element; tbl_element.getparent().remove(tbl_element)
            logger.info(f"{log_prefix} Conteúdo principal do template limpo.")

            # Verifica a existência do estilo 'Normal' UMA VEZ
            if NORMAL_STYLE_NAME in doc.styles:
                normal_style_exists = True
                logger.info(f"{log_prefix} Estilo '{NORMAL_STYLE_NAME}' encontrado no template.")
            else:
                logger.warning(f"{log_prefix} AVISO: Estilo '{NORMAL_STYLE_NAME}' NÃO encontrado no template '{os.path.basename(template_docx_path)}'. A formatação de parágrafos normais dependerá do padrão do documento.")

        except Exception as e_load_template:
             logger.error(f"{log_prefix} FATAL: Falha ao carregar/limpar template OBRIGATÓRIO '{os.path.basename(template_docx_path)}': {e_load_template}. Abortando.")
             return (False, None)

    except Exception as e_doc:
        logger.error(f"{log_prefix} FATAL: Erro crítico ao preparar DOCX: {e_doc}")
        logger.error(traceback.format_exc()); return (False, None)

    # Processa os chunks pela API e formata o DOCX
    logger.info(f"{log_prefix} Iniciando chamadas à API para CORREÇÃO de {len(text_chunks)} chunks...")
    corrected_text_list_pass1 = [] # Acumula texto corrigido para o Passo 2
    processed_chunks_count = 0
    failed_chunks_count = 0
    # Barra de progresso
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P1: Corrigindo", unit="chunk", leave=False)

    for i, chunk in progress_bar:
        # Chama a IA para correção
        corrected_chunk = format_with_ai_correction_only(model, chunk, author_name, book_name, is_first_chunk=(i == 0))

        # Aplica formatação ao DOCX e acumula texto
        if corrected_chunk is not None: # Checa explicitamente por None (falha API)
            # Processa mesmo se string vazia (pode ser intencional da IA)
             try:
                apply_formatting_pass1(doc, corrected_chunk, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
                processed_chunks_count += 1
             except Exception as format_err:
                logger.error(f"{log_prefix} Erro na apply_formatting_pass1 (Chunk {i+1}): {format_err}.")
                logger.error(traceback.format_exc())
                failed_chunks_count += 1
                # Aplica fallback com marcador de erro de formatação
                try:
                    fallback_text = f"{FORMATTING_ERROR_MARKER}\n\n{chunk}"
                    apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
                except Exception as fallback_format_err:
                    logger.critical(f"{log_prefix} Falha CRÍTICA ao aplicar fallback de erro FORMATAÇÃO (Chunk {i+1}): {fallback_format_err}.")
        else: # Falha na chamada da API (retornou None)
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na CORREÇÃO (API retornou None). Usando fallback com marcador.")
            failed_chunks_count += 1
            # Aplica fallback com marcador de falha da API
            try:
                fallback_text = f"{AI_FAILURE_MARKER}\n\n{chunk}"
                apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
            except Exception as fallback_format_err:
                 logger.critical(f"{log_prefix} Falha CRÍTICA ao aplicar fallback de falha API (Chunk {i+1}): {fallback_format_err}.")

        # --- Salva progresso DOCX periodicamente ---
        processed_total = processed_chunks_count + failed_chunks_count
        # Salva a cada 10 chunks processados ou no último chunk
        if processed_total > 0 and (processed_total % 10 == 0 or (i + 1) == len(text_chunks)):
            temp_save_path = f"{output_docx_path}.{processed_total}.temp_save" # Nome temporário
            try:
                # logger.debug(f"{log_prefix} Salvando progresso DOCX (chunk {i+1})...") # Verbose
                doc.save(temp_save_path)
                # Move atômico é geralmente mais seguro
                shutil.move(temp_save_path, output_docx_path)
                logger.info(f"{log_prefix} Progresso DOCX (Passo 1) salvo ({processed_total} chunks processados).")
            except Exception as e_save:
                 logger.error(f"{log_prefix} Erro ao salvar progresso DOCX (Chunk {i+1}) para '{os.path.basename(output_docx_path)}': {e_save}")
                 # Tenta remover arquivo temporário possivelmente corrompido
                 if os.path.exists(temp_save_path):
                      try: os.remove(temp_save_path)
                      except OSError: pass

    # --- Salva o DOCX final após o loop ---
    try:
        logger.info(f"{log_prefix} Salvando DOCX final (Passo 1) em: {os.path.basename(output_docx_path)}")
        doc.save(output_docx_path)
    except Exception as e_final_save:
        logger.error(f"{log_prefix} Erro no salvamento final do DOCX (Passo 1): {e_final_save}")
        logger.error(traceback.format_exc())
        # Continua para retornar o texto, mas DOCX pode estar incompleto/corrompido

    # Junta o texto corrigido para retornar ao Passo 2
    full_corrected_text = "\n\n".join(corrected_text_list_pass1)
    logger.info(f"{log_prefix} Acumulado texto corrigido para Pass 2 ({len(full_corrected_text)} chars).")

    logger.info(f"{log_prefix} --- Passo 1 concluído. Chunks OK: {processed_chunks_count}, Falhas/Fallback: {failed_chunks_count} ---")
    # Retorna sucesso (True mesmo se alguns chunks falharam mas DOCX foi salvo) e o texto
    return (True, full_corrected_text)


def run_footnote_id_pass(model, corrected_text_content, author_name, book_name):
    """
    Executa o Passo 2: Identifica notas no texto já corrigido pela IA.
    Retorna: (bool, str | None) -> (success_status, marked_text_content)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 2: Identificação de Notas ---")
    if corrected_text_content is None: # Checa explicitamente por None
        logger.error(f"{log_prefix} Texto corrigido de entrada é None. Abortando Passo 2.")
        return (False, None)
    # logger.info(f"{log_prefix} Recebido texto corrigido para Passo 2 ({len(corrected_text_content)} chars).") # Verbose

    # Divide o texto corrigido em chunks
    logger.info(f"{log_prefix} Dividindo texto corrigido em chunks para notas...")
    text_chunks = create_chunks(corrected_text_content, MAX_CHUNK_TOKENS, author_name, book_name)
    if not text_chunks:
        logger.error(f"{log_prefix} Nenhum chunk gerado do texto corrigido. Abortando Passo 2."); return (False, None)
    # Removido log incorreto que dizia 1 chunk
    # logger.info(f"{log_prefix} Texto corrigido dividido em {len(text_chunks)} chunks.") # Correto

    # Processa chunks pela API para marcar notas
    logger.info(f"{log_prefix} Iniciando chamadas à API para IDENTIFICAÇÃO DE NOTAS em {len(text_chunks)} chunks...")
    marked_text_list_pass2 = []
    processed_chunks_count = 0
    failed_chunks_count = 0
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P2: Notas", unit="chunk", leave=False)

    for i, chunk in progress_bar:
        # Chama a IA para marcar notas
        marked_chunk = format_with_ai_footnote_only(model, chunk, author_name, book_name)

        if marked_chunk is not None: # API retornou algo (pode ou não ter marcadores)
            marked_text_list_pass2.append(marked_chunk)
            processed_chunks_count += 1
            # if "[NOTA_" in marked_chunk: logger.debug(f"{log_prefix} Chunk {i+1}: marcadores de nota encontrados.") # Verbose
            # else: logger.debug(f"{log_prefix} Chunk {i+1}: NENHUM marcador de nota adicionado.") # Verbose
        else: # Falha na API
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na IDENTIFICAÇÃO DE NOTAS (API retornou None). Usando texto original do chunk.")
            marked_text_list_pass2.append(chunk) # Usa o chunk original como fallback
            failed_chunks_count += 1

    # Junta o texto marcado para retornar ao Passo 3
    full_marked_text = "\n\n".join(marked_text_list_pass2)
    logger.info(f"{log_prefix} Acumulado texto com marcadores para Pass 3 ({len(full_marked_text)} chars).")

    logger.info(f"{log_prefix} --- Passo 2 concluído. Chunks OK: {processed_chunks_count}, Falhas/Fallback: {failed_chunks_count} ---")
    # Retorna sucesso e o texto marcado
    return (True, full_marked_text)


def run_final_txt_generation(marked_text_content, output_notes_path, output_numbered_txt_path, author_name, book_name):
    """
    Executa o Passo 3: Processa marcadores [NOTA_...] para gerar TXT final numerado [N]
            e um arquivo separado com as notas.
    Retorna: bool -> success_status
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 3: Geração Final TXT (Notas e Numerado) ---")
    if marked_text_content is None: # Checa explicitamente por None
        logger.error(f"{log_prefix} Texto marcado de entrada é None. Abortando Passo 3.")
        return False
    # logger.info(f"{log_prefix} Recebido texto com marcadores para Passo 3 ({len(marked_text_content)} chars).") # Verbose

    # Garante que diretórios de saída existem
    os.makedirs(os.path.dirname(output_notes_path), exist_ok=True)
    os.makedirs(os.path.dirname(output_numbered_txt_path), exist_ok=True)

    footnote_counter = 1
    notes_found = []
    # Regex para encontrar pares de marcadores: [NOTA_TIPO:Ref][CONTEUDO_NOTA:Conteúdo]
    footnote_pattern = re.compile(
        r'(\[NOTA_(?:IDIOMA|CITACAO|NOME|TERMO):[^\]]+?\])\s*(\[CONTEUDO_NOTA:([^\]]*?)\])',
        re.IGNORECASE # Ignora maiúsculas/minúsculas nos marcadores
    )

    # Função para substituir o marcador pelo número e coletar a nota
    def replace_marker_and_collect_note(match):
        nonlocal footnote_counter
        original_marker = match.group(1) # Ex: [NOTA_NOME:Kropotkin]
        content_marker = match.group(2)  # Ex: [CONTEUDO_NOTA:Piotr Kropotkin...]
        content = match.group(3).strip() # Ex: Piotr Kropotkin...

        # Lida com casos onde CONTEUDO_NOTA está vazio
        if not content:
             logger.warning(f"{log_prefix} Encontrado marcador [CONTEUDO_NOTA:] vazio após {original_marker}. Ignorando nota.")
             return "" # Remove ambos os marcadores sem adicionar número

        notes_found.append(f"{footnote_counter}. {content}") # Armazena a nota numerada
        replacement = f"[{footnote_counter}]" # Texto que substituirá os marcadores no texto principal
        # logger.debug(f"{log_prefix} Nota {footnote_counter}: '{content}'") # Verbose
        footnote_counter += 1
        return replacement

    logger.info(f"{log_prefix} Processando marcadores e gerando arquivos finais TXT...")
    try:
        # Processa o texto inteiro de uma vez usando a função de substituição
        final_numbered_text = footnote_pattern.sub(replace_marker_and_collect_note, marked_text_content)

        # --- Salva o arquivo de notas ---
        logger.info(f"{log_prefix} Salvando arquivo de notas em: {os.path.basename(output_notes_path)}")
        with open(output_notes_path, "w", encoding="utf-8") as f_notes:
            f_notes.write(f"Notas de Rodapé Geradas para {author_name} - {book_name}\n")
            f_notes.write("=" * (30 + len(author_name) + len(book_name)) + "\n\n")
            if notes_found:
                f_notes.write("\n".join(notes_found))
                f_notes.write("\n") # Linha extra no final
                logger.info(f"{log_prefix} {len(notes_found)} notas salvas em '{os.path.basename(output_notes_path)}'.")
            else:
                f_notes.write("(Nenhuma nota de rodapé foi identificada ou extraída com sucesso)\n")
                logger.info(f"{log_prefix} Nenhuma nota de rodapé identificada/salva.")

        # --- Salva o TXT final com as referências numeradas [N] ---
        logger.info(f"{log_prefix} Salvando TXT final com números [{footnote_counter-1}] em: {os.path.basename(output_numbered_txt_path)}")
        with open(output_numbered_txt_path, "w", encoding="utf-8") as f_numbered:
            f_numbered.write(final_numbered_text)
        logger.info(f"{log_prefix} TXT final com números salvo ({len(final_numbered_text)} chars).")

    except Exception as e_final_gen:
        logger.error(f"{log_prefix} FATAL: Erro durante a geração final dos arquivos TXT (Passo 3): {e_final_gen}")
        logger.error(traceback.format_exc())
        return False # Indica falha

    logger.info(f"{log_prefix} --- Passo 3 concluído. ---")
    return True # Indica sucesso

# --- Funções para Gerenciar Arquivos Processados (CORREÇÃO) ---

def load_processed_files(filepath):
    """Lê o arquivo de log de CORREÇÃO e retorna um set."""
    processed = set()
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                cleaned_line = line.strip()
                if cleaned_line:
                    processed.add(cleaned_line)
        logger.info(f"Carregados {len(processed)} registros de CORREÇÕES concluídas de '{filepath}'.")
    except FileNotFoundError:
        logger.info(f"Arquivo de log de correções '{filepath}' não encontrado. Iniciando sem histórico.")
    except Exception as e:
        logger.error(f"Erro ao carregar log de correções '{filepath}': {e}")
    return processed

def log_processed_file(filepath, file_identifier):
    """Adiciona um identificador de arquivo ao log de CORREÇÃO."""
    try:
        with open(filepath, 'a', encoding='utf-8') as f:
            f.write(f"{file_identifier}\n")
        # logger.debug(f"Registrado '{file_identifier}' como CORRIGIDO em '{filepath}'.") # Verbose
    except Exception as e:
        logger.error(f"Erro ao registrar '{file_identifier}' no log de correções '{filepath}': {e}")

# --- Funções para Gerenciar Arquivos Processados (TRADUÇÃO) --- << NOVAS >>

def load_translated_files(filepath):
    """Lê o arquivo de log de TRADUÇÃO e retorna um set."""
    processed = set()
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                cleaned_line = line.strip()
                if cleaned_line:
                    processed.add(cleaned_line)
        logger.info(f"Carregados {len(processed)} registros de TRADUÇÕES concluídas de '{filepath}'.")
    except FileNotFoundError:
        logger.info(f"Arquivo de log de traduções '{filepath}' não encontrado. Iniciando sem histórico de traduções.")
    except Exception as e:
        logger.error(f"Erro ao carregar log de traduções '{filepath}': {e}")
    return processed

def log_translated_file(filepath, file_identifier):
    """Adiciona um identificador de arquivo ao log de TRADUÇÃO."""
    try:
        with open(filepath, 'a', encoding='utf-8') as f:
            f.write(f"{file_identifier}\n")
        logger.debug(f"Registrado '{file_identifier}' como TRADUZIDO com sucesso em '{filepath}'.") # Verbose
    except Exception as e:
        logger.error(f"Erro ao registrar '{file_identifier}' no log de traduções '{filepath}': {e}")

# --- FUNÇÃO DE ENVIO DE E-MAIL (MODIFICADA) ---
def send_completion_email(sender_email, sender_password, recipient_email, smtp_server, smtp_port,
                          processed_correction, skipped_correction, failed_correction, # Contadores Correção
                          processed_translation, skipped_translation, failed_translation, # Contadores Tradução (NOVOS)
                          total_duration_seconds,
                          main_log_path, processed_log_path, translated_log_path): # Caminho do log de tradução (NOVO)
    """Envia um e-mail de notificação de conclusão com resumo de ambas as etapas."""

    if not sender_email or not sender_password or not recipient_email: # Adicionado recipient_email na checagem
        logger.warning("E-mail de envio, senha ou destinatário não configurados no .env ou código. Não é possível enviar notificação.")
        return

    logger.info(f"Preparando e-mail de notificação para {recipient_email}...")

    subject = "Script Processador de Livros (Correção + Tradução) - Conclusão" # Assunto atualizado
    body = f"""
Olá,

O script de processamento de livros (Correção + Tradução Híbrida) concluiu a execução.

Resumo Geral:
--------------------------------------------------
- Tempo Total de Execução: {total_duration_seconds:.2f} segundos ({total_duration_seconds/60:.2f} minutos)

Resumo Etapa de Correção:
--------------------------------------------------
- Livros Corrigidos com Sucesso (nesta execução): {processed_correction}
- Livros Pulados (correção já feita antes): {skipped_correction}
- Livros com Falha na Correção: {failed_correction}

Resumo Etapa de Tradução Híbrida:
--------------------------------------------------
- Livros Traduzidos com Sucesso (nesta execução): {processed_translation}
- Livros Pulados (tradução já feita antes): {skipped_translation}
- Livros com Falha na Tradução (ou input ausente): {failed_translation}
--------------------------------------------------

Logs para Consulta:
- Log Detalhado da Execução: {os.path.abspath(main_log_path)}
- Log de Correções Concluídas: {os.path.abspath(processed_log_path)}
- Log de Traduções Concluídas: {os.path.abspath(translated_log_path)}

Atenciosamente,
Seu Script Processador de Livros
"""

    message = EmailMessage()
    message['Subject'] = subject
    message['From'] = sender_email
    message['To'] = recipient_email
    message.set_content(body)

    context = ssl.create_default_context() # For secure connection

    try:
        server = None # Initialize server variable
        logger.info(f"Conectando ao servidor SMTP: {smtp_server}:{smtp_port}...")
        if smtp_port == 465: # SSL connection
             server = smtplib.SMTP_SSL(smtp_server, smtp_port, context=context, timeout=30) # Timeout adicionado
             server.login(sender_email, sender_password)
        else: # Assume TLS (port 587 or other)
            server = smtplib.SMTP(smtp_server, smtp_port, timeout=30) # Timeout adicionado
            server.ehlo() # Identify client to server
            server.starttls(context=context) # Secure the connection
            server.ehlo() # Re-identify after TLS
            server.login(sender_email, sender_password)

        logger.info("Enviando e-mail de resumo final...")
        server.send_message(message)
        logger.info(f"✅ E-mail de resumo final enviado com sucesso para {recipient_email}.")

    except smtplib.SMTPAuthenticationError:
        logger.error("FALHA NA AUTENTICAÇÃO do e-mail. Verifique EMAIL_SENDER_ADDRESS e EMAIL_SENDER_APP_PASSWORD no .env.")
        logger.error("Lembre-se: Para Gmail com 2FA, use uma 'Senha de App'.")
    except smtplib.SMTPServerDisconnected:
         logger.error("Servidor SMTP desconectou inesperadamente. Tente novamente.")
    except smtplib.SMTPConnectError as e:
         logger.error(f"Erro ao conectar ao servidor SMTP {smtp_server}:{smtp_port}. Verifique o nome/porta e a rede. Erro: {e}")
    except smtplib.SMTPException as e:
        logger.error(f"Erro SMTP ao enviar e-mail: {e}")
        logger.error(traceback.format_exc())
    except ssl.SSLError as e:
         logger.error(f"Erro SSL/TLS ao conectar ao servidor SMTP: {e}")
         logger.error("Verifique a porta e as configurações de segurança (SSL/TLS).")
    except OSError as e:
         logger.error(f"Erro de Rede/OS (ex: Timeout, Host não encontrado) ao tentar conectar ao servidor SMTP: {e}")
         logger.error("Verifique a conexão com a internet e o endereço/porta do servidor.")
    except Exception as e:
        logger.error(f"Erro inesperado ao enviar e-mail: {e}")
        logger.error(traceback.format_exc())
    finally:
        if server:
            try:
                server.quit() # Ensure connection is closed
            except smtplib.SMTPServerDisconnected:
                 pass # Ignore if already disconnected
            except Exception as e_quit:
                 logger.warning(f"Erro ao fechar conexão SMTP: {e_quit}")


# --- FUNÇÃO PRINCIPAL (main) ---
def main():
    start_time_main = time.time()
    logger.info("========================================================")
    logger.info(f"Iniciando Processador Multi-Autor/Livro (Correção + Tradução Híbrida) - {time.strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info(f"Diretório de Entrada TXT: {BASE_INPUT_TXT_DIR}")
    logger.info(f"Diretório de Saída DOCX: {BASE_OUTPUT_DOCX_DIR}")
    logger.info(f"Diretório de Saída TXT (Notas/Final): {BASE_OUTPUT_TXT_DIR}")
    logger.info(f"Template DOCX OBRIGATÓRIO: {TEMPLATE_DOCX}")
    logger.info(f"Log de Correções Concluídas: {PROCESSED_LOG_FILE}")
    logger.info(f"Log de Traduções Concluídas: {TRANSLATED_LOG_FILE}")
    logger.info(f"Script Tradutor a ser chamado: {PATH_TO_TRANSLATOR_SCRIPT}")
    logger.info(f"Número de palavras para traduzir: {NUM_WORDS_TO_TRANSLATE}")
    logger.info("========================================================")

    # Carrega a lista de arquivos já processados (CORREÇÃO)
    processed_files_set = load_processed_files(PROCESSED_LOG_FILE)
    # Carrega a lista de arquivos já processados (TRADUÇÃO)
    translated_files_set = load_translated_files(TRANSLATED_LOG_FILE)

    # --- Valida Diretório de Entrada ---
    if not os.path.isdir(BASE_INPUT_TXT_DIR):
        logger.error(f"FATAL: Diretório de entrada base '{BASE_INPUT_TXT_DIR}' não encontrado! Abortando.")
        return

    # --- Encontra Pastas de Autores ---
    try:
        author_folders = sorted([f for f in os.listdir(BASE_INPUT_TXT_DIR) if os.path.isdir(os.path.join(BASE_INPUT_TXT_DIR, f))])
    except Exception as e:
        logger.error(f"FATAL: Erro ao listar diretórios de autores em '{BASE_INPUT_TXT_DIR}': {e}")
        return
    if not author_folders:
        logger.warning(f"Nenhuma subpasta de autor encontrada em '{BASE_INPUT_TXT_DIR}'. Saindo.")
        return
    logger.info(f"Autores encontrados ({len(author_folders)}): {', '.join(author_folders)}")

    # --- Inicializa Contadores para Resumo Final ---
    total_books_processed_correction = 0 # Sucesso na correção nesta run
    total_books_skipped_correction = 0   # Correção já feita antes
    total_books_failed_correction = 0    # Falha na correção nesta run
    total_translation_processed = 0      # Sucesso na tradução nesta run
    total_translation_skipped = 0        # Tradução já feita antes
    total_translation_failed = 0         # Falha na tradução nesta run (ou input ausente)

    # === LOOP PRINCIPAL: Itera por cada pasta de AUTOR ===
    for author_name in author_folders:
        author_input_dir = os.path.join(BASE_INPUT_TXT_DIR, author_name)
        author_output_docx_dir = os.path.join(BASE_OUTPUT_DOCX_DIR, author_name)
        author_output_txt_dir = os.path.join(BASE_OUTPUT_TXT_DIR, author_name)

        logger.info(f"--- Verificando Autor: {author_name} em '{author_input_dir}' ---")

        # --- Encontra arquivos .txt de LIVROS dentro da pasta do autor ---
        try:
            input_txt_files = sorted(glob.glob(os.path.join(author_input_dir, "*.txt")))
            input_txt_files_filtered = [
                f for f in input_txt_files if not (
                    os.path.basename(f).endswith(FINAL_NUMBERED_TXT_BASENAME) or
                    os.path.basename(f).endswith(NOTES_TXT_FILE_BASENAME) or
                    os.path.basename(f).startswith("backup_")
                )
            ]
            input_txt_files = input_txt_files_filtered
        except Exception as e:
            logger.error(f"[{author_name}] Erro ao listar/filtrar arquivos .txt em '{author_input_dir}': {e}. Pulando autor.")
            continue

        if not input_txt_files:
            logger.warning(f"[{author_name}] Nenhum arquivo .txt de entrada (livro) encontrado/restante em '{author_input_dir}'.")
            continue

        logger.info(f"[{author_name}] Encontrados {len(input_txt_files)} arquivos .txt potenciais para processar.")

        # === LOOP INTERNO: Itera por cada LIVRO (arquivo .txt) ===
        for input_txt_path in input_txt_files:
            book_filename = os.path.basename(input_txt_path)
            file_identifier = f"{author_name}/{book_filename}"
            log_prefix_book = f"[{file_identifier}]"

            logger.info(f"--------------------------------------------------------")
            logger.info(f"{log_prefix_book} Processando Livro...")

            correction_successful_this_run = False
            output_docx_path = None
            base_book_name = os.path.splitext(book_filename)[0]
            author_output_docx_dir = os.path.join(BASE_OUTPUT_DOCX_DIR, author_name)

            # --- Verifica Status da CORREÇÃO ---
            logger.info(f"{log_prefix_book} Verificando status (Correção)...")
            if file_identifier in processed_files_set:
                logger.info(f"{log_prefix_book} CORREÇÃO já realizada anteriormente (encontrado em '{PROCESSED_LOG_FILE}'). Pulando etapa de correção.")
                total_books_skipped_correction += 1
                correction_successful_this_run = True
                output_docx_path = os.path.join(author_output_docx_dir, f"{base_book_name}_{FINAL_DOCX_BASENAME}")
            else:
                # --- Executa a CORREÇÃO ---
                logger.info(f"{log_prefix_book} Iniciando processamento (Correção)...")
                book_start_time = time.time()

                output_docx_path = os.path.join(author_output_docx_dir, f"{base_book_name}_{FINAL_DOCX_BASENAME}")
                output_notes_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{NOTES_TXT_FILE_BASENAME}")
                output_numbered_txt_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{FINAL_NUMBERED_TXT_BASENAME}")

                try:
                    os.makedirs(author_output_docx_dir, exist_ok=True)
                    os.makedirs(author_output_txt_dir, exist_ok=True)
                except Exception as e_mkdir:
                     logger.error(f"{log_prefix_book} ERRO ao criar diretórios de saída: {e_mkdir}. Pulando este livro.")
                     total_books_failed_correction += 1
                     continue

                all_correction_steps_successful = True
                corrected_text_content = None
                marked_text_content = None

                try:
                    # PASSO 1: CORREÇÃO DOCX
                    pass1_success, corrected_text_content = run_correction_pass(
                        gemini_model, input_txt_path, TEMPLATE_DOCX, output_docx_path,
                        author_name, base_book_name
                    )
                    if not pass1_success or corrected_text_content is None:
                        logger.error(f"{log_prefix_book} Passo 1 (Correção/DOCX) FALHOU ou retornou vazio.")
                        all_correction_steps_successful = False
                    else:
                        # PASSO 2: IDENTIFICAÇÃO DE NOTAS
                        pass2_success, marked_text_content = run_footnote_id_pass(
                            gemini_model, corrected_text_content,
                            author_name, base_book_name
                        )
                        if not pass2_success or marked_text_content is None:
                            logger.error(f"{log_prefix_book} Passo 2 (Identificação Notas) FALHOU ou retornou vazio.")
                            all_correction_steps_successful = False
                        else:
                            # PASSO 3: GERAÇÃO FINAL TXT
                            pass3_success = run_final_txt_generation(
                                marked_text_content, output_notes_path, output_numbered_txt_path,
                                author_name, base_book_name
                            )
                            if not pass3_success:
                                logger.error(f"{log_prefix_book} Passo 3 (Geração TXTs Finais) FALHOU.")
                                all_correction_steps_successful = False
                except Exception as e_corr_steps:
                     logger.error(f"{log_prefix_book} Erro inesperado durante os passos de CORREÇÃO: {e_corr_steps}")
                     logger.error(traceback.format_exc())
                     all_correction_steps_successful = False

                book_end_time = time.time()
                book_total_time = book_end_time - book_start_time

                if all_correction_steps_successful:
                    logger.info(f"✅ {log_prefix_book} Etapa de CORREÇÃO concluída com SUCESSO em {book_total_time:.2f} seg.")
                    log_processed_file(PROCESSED_LOG_FILE, file_identifier)
                    processed_files_set.add(file_identifier)
                    total_books_processed_correction += 1
                    correction_successful_this_run = True
                else:
                    logger.warning(f"⚠️ {log_prefix_book} Etapa de CORREÇÃO concluída com FALHAS em {book_total_time:.2f} seg. NÃO será traduzido nesta execução.")
                    total_books_failed_correction += 1
                    correction_successful_this_run = False

            # --- ETAPA DE TRADUÇÃO (Só executa se a CORREÇÃO funcionou) ---
            if correction_successful_this_run:
                logger.info(f"{log_prefix_book} Verificando status (Tradução)...")

                if output_docx_path is None or not os.path.exists(output_docx_path):
                     logger.warning(f"{log_prefix_book} Arquivo DOCX de entrada para tradução não encontrado ('{os.path.basename(output_docx_path or 'N/A')}'). Pulando tradução.")
                     total_translation_failed += 1
                elif file_identifier in translated_files_set:
                    logger.info(f"{log_prefix_book} TRADUÇÃO já realizada anteriormente (encontrado em '{TRANSLATED_LOG_FILE}'). Pulando.")
                    total_translation_skipped += 1
                else:
                    logger.info(f"{log_prefix_book} >>> Iniciando etapa de TRADUÇÃO HÍBRIDA...")
                    translation_start_time = time.time()
                    translation_successful = False

                    translated_docx_path = os.path.join(author_output_docx_dir, f"{base_book_name}_{FINAL_DOCX_BASENAME.replace('.docx', TRANSLATED_DOCX_SUFFIX)}")

                    if not os.path.exists(PATH_TO_TRANSLATOR_SCRIPT):
                         logger.error(f"{log_prefix_book} ERRO CRÍTICO: Script tradutor não encontrado em '{PATH_TO_TRANSLATOR_SCRIPT}'. Verifique a constante no script. Pulando tradução.")
                         total_translation_failed += 1
                    else:
                        try:
                            command = [
                                sys.executable, # Usa o python do venv
                                PATH_TO_TRANSLATOR_SCRIPT,
                                '--input', output_docx_path,
                                '--output', translated_docx_path,
                                '--words', str(NUM_WORDS_TO_TRANSLATE)
                            ]
                            logger.info(f"{log_prefix_book} Executando comando: {' '.join(command)}")

                            result = subprocess.run(command, capture_output=True, text=True, encoding='utf-8', check=False)

                            translation_end_time = time.time()
                            translation_total_time = translation_end_time - translation_start_time

                            if result.returncode == 0:
                                logger.info(f"✅ {log_prefix_book} Etapa de TRADUÇÃO HÍBRIDA concluída com SUCESSO em {translation_total_time:.2f} seg.")
                                log_translated_file(TRANSLATED_LOG_FILE, file_identifier)
                                translated_files_set.add(file_identifier)
                                total_translation_processed += 1
                                translation_successful = True
                                if result.stdout: logger.debug(f"{log_prefix_book} Saída do script tradutor:\n{result.stdout}")
                            else:
                                logger.error(f"❌ {log_prefix_book} Etapa de TRADUÇÃO HÍBRIDA FALHOU (script retornou código: {result.returncode}) em {translation_total_time:.2f} seg.")
                                if result.stderr: logger.error(f"{log_prefix_book} Erro reportado pelo script tradutor:\n{result.stderr}")
                                else: logger.error(f"{log_prefix_book} Script tradutor não retornou mensagem de erro específica no stderr.")
                                total_translation_failed += 1

                        except FileNotFoundError:
                             logger.error(f"{log_prefix_book} ERRO: Comando '{sys.executable}' não encontrado ou script '{PATH_TO_TRANSLATOR_SCRIPT}' não encontrado pelo sistema. Verifique o PATH e o caminho do script. Pulando tradução.")
                             total_translation_failed += 1
                        except Exception as e_translate_sub:
                             logger.error(f"{log_prefix_book} Erro CRÍTICO ao tentar executar o subprocesso de tradução: {e_translate_sub}")
                             logger.error(traceback.format_exc())
                             total_translation_failed += 1
            # --- Fim da verificação/execução da tradução ---

            logger.info(f"{log_prefix_book} --- Fim do processamento do livro ---")

        # Fim do loop de livros para o autor atual
        logger.info(f"--- Concluída verificação do Autor: {author_name} ---")
        logger.info(f"--------------------------------------------------------")


    # --- Fim do Loop Principal (Todos os Autores) ---
    end_time_main = time.time()
    total_time_main = end_time_main - start_time_main

    # === Resumo Final Logging ===
    logger.info("===================== RESUMO FINAL =====================")
    logger.info(f"Tempo total geral de execução: {total_time_main:.2f} seg ({total_time_main/60:.2f} min).")
    logger.info("--- Resumo Etapa de Correção ---")
    logger.info(f"Livros Corrigidos com Sucesso (nesta execução): {total_books_processed_correction}")
    logger.info(f"Livros Pulados (correção já feita anteriormente): {total_books_skipped_correction}")
    logger.info(f"Livros com Falha na Correção (nesta execução): {total_books_failed_correction}")
    logger.info("--- Resumo Etapa de Tradução Híbrida ---")
    logger.info(f"Livros Traduzidos com Sucesso (nesta execução): {total_translation_processed}")
    logger.info(f"Livros Pulados (tradução já feita anteriormente): {total_translation_skipped}")
    logger.info(f"Livros com Falha na Tradução (nesta execução): {total_translation_failed}")
    logger.info("--- Logs ---")
    logger.info(f"Log detalhado: {os.path.abspath(log_filepath)}")
    logger.info(f"Log de correções concluídas: {os.path.abspath(PROCESSED_LOG_FILE)}")
    logger.info(f"Log de traduções concluídas: {os.path.abspath(TRANSLATED_LOG_FILE)}")
    logger.info("--- Arquivos Gerados (Exemplos) ---")
    logger.info(f"  - DOCX Corrigido: {BASE_OUTPUT_DOCX_DIR}/<Autor>/<Livro>_{FINAL_DOCX_BASENAME}")
    logger.info(f"  - TXT Numerado:   {BASE_OUTPUT_TXT_DIR}/<Autor>/<Livro>_{FINAL_NUMBERED_TXT_BASENAME}")
    logger.info(f"  - TXT Notas:      {BASE_OUTPUT_TXT_DIR}/<Autor>/<Livro>_{NOTES_TXT_FILE_BASENAME}")
    logger.info(f"  - DOCX Traduzido: {BASE_OUTPUT_DOCX_DIR}/<Autor>/<Livro>_{FINAL_DOCX_BASENAME.replace('.docx', TRANSLATED_DOCX_SUFFIX)}")
    logger.info("========================================================")

    # === Envio de E-mail FINAL (MODIFICADO para incluir stats de tradução) ===
    if EMAIL_SENDER_ADDRESS and EMAIL_SENDER_APP_PASSWORD and EMAIL_RECIPIENT_ADDRESS:
        # Chama a função de e-mail passando TODOS os contadores relevantes
        send_completion_email(
            sender_email=EMAIL_SENDER_ADDRESS,
            sender_password=EMAIL_SENDER_APP_PASSWORD,
            recipient_email=EMAIL_RECIPIENT_ADDRESS,
            smtp_server=EMAIL_SMTP_SERVER,
            smtp_port=EMAIL_SMTP_PORT,
            # Passa contadores da CORREÇÃO
            processed_correction=total_books_processed_correction,
            skipped_correction=total_books_skipped_correction,
            failed_correction=total_books_failed_correction,
            # Passa contadores da TRADUÇÃO (NOVOS)
            processed_translation=total_translation_processed,
            skipped_translation=total_translation_skipped,
            failed_translation=total_translation_failed,
            # Passa duração e logs
            total_duration_seconds=total_time_main,
            main_log_path=log_filepath,
            processed_log_path=PROCESSED_LOG_FILE,
            translated_log_path=TRANSLATED_LOG_FILE # Passa o novo log também
        )
    else:
        logger.info("Envio de e-mail de resumo final pulado (configuração ausente ou incompleta no .env).")


# --- Ponto de Entrada do Script ---
if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("\nProcesso interrompido manualmente (Ctrl+C). Status final pode estar incompleto.")
    except Exception as e_main:
        logger.critical(f"FATAL: Erro inesperado e não tratado na execução de main(): {e_main}")
        logger.critical(traceback.format_exc())
        # Considerar enviar um e-mail de falha crítica aqui, se possível